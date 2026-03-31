import docx
import re

def extract_and_analyze_markers(file_path):
    doc = docx.Document(file_path)
    if not doc.tables:
        print("No tables found.")
        return

    table = doc.tables[0]
    if len(table.rows) < 2:
        print("Table too small.")
        return

    content_row = table.rows[1]
    
    def get_markers(paragraphs):
        markers = []
        for p in paragraphs:
            text = p.text.strip()
            # Match markers like { 1023 } or {1023}
            match = re.search(r'\{?\s*\d+\s*\}?', text)
            if match:
                markers.append(match.group(0))
        return markers

    en_markers = get_markers(content_row.cells[0].paragraphs)
    ru_markers = get_markers(content_row.cells[1].paragraphs)
    uz_markers = get_markers(content_row.cells[2].paragraphs)

    print(f"EN markers count: {len(en_markers)}")
    print(f"RU markers count: {len(ru_markers)}")
    print(f"UZ markers count: {len(uz_markers)}")
    
    print("\nFirst 10 EN markers:", en_markers[:10])
    print("First 10 RU markers:", ru_markers[:10])
    print("First 10 UZ markers:", uz_markers[:10])

if __name__ == "__main__":
    file_path = r"c:\Users\Администратор\Desktop\2\1023. рус авто v2.docx"
    extract_and_analyze_markers(file_path)
