"""
Import 'Таҳририят кенгаши тасдиқлаган мунозарали сўзлар' from docx.
Source: Qonun/SP/Умумий Мунозарали бўлган сўзлар.docx (234 rows, 8 cols)

Columns in docx:
  0. Т/р (No)
  1. Таржима қиланаётган сўзнинг рус тилидаги номланиши
  2. Европа фармакопеясида келтирилган матн (EN)
  3. Европа фармакопеясининг рус тилидаги таржима матни
  4. Изоҳли тарифи
  5. Мавжуд вариантлар
  6. Таклиф этилаётган вариант
  7. Адабиётлар

Target table: disputed_board
"""
import os
import sys
import sqlite3
import logging
from pathlib import Path

logging.basicConfig(level=logging.INFO, format="[disputed_board] %(message)s")
log = logging.getLogger()

DB_PATH = os.getenv("DB_PATH", os.path.join(os.path.dirname(os.path.dirname(__file__)), "pharma_editor.db"))
DATA_DIR = Path(os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "source_data", "disputed_board"
))


def ensure_schema():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS disputed_board (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            seq_no INTEGER,
            ru_term TEXT,
            en_context TEXT,
            ru_context TEXT,
            definition_uz TEXT,
            existing_variants TEXT,
            proposed_variant TEXT,
            references_text TEXT,
            source TEXT DEFAULT 'editorial_board',
            status TEXT DEFAULT 'approved',
            modified_by TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_db_ru ON disputed_board(ru_term)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_db_proposed ON disputed_board(proposed_variant)")
    conn.commit()
    conn.close()


def import_docx(data_dir: Path = None) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed"}

    d = data_dir or DATA_DIR
    p = d / "disputed_board.docx"
    if not p.exists():
        log.warning(f"File not found: {p}")
        return {"error": "file not found", "path": str(p)}

    ensure_schema()
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    # No idempotent guard — INSERT OR IGNORE handles duplicates
    doc = Document(str(p))
    if not doc.tables:
        conn.close()
        return {"error": "no tables in docx"}

    table = doc.tables[0]
    added = 0
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            continue  # header
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 8:
            continue
        try:
            seq = int(cells[0]) if cells[0].isdigit() else (r_idx)
        except Exception:
            seq = r_idx

        ru_term = cells[1]
        if not ru_term:
            continue
        try:
            # Write to legacy disputed_board table
            cur.execute("""
                INSERT OR IGNORE INTO disputed_board
                (seq_no, ru_term, en_context, ru_context, definition_uz,
                 existing_variants, proposed_variant, references_text)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (seq, ru_term, cells[2], cells[3], cells[4], cells[5], cells[6], cells[7]))
            # Also write to UI-facing disputed_words table
            cur.execute("""
                INSERT INTO disputed_words
                (en, ru, uz, context_en, context_ru, context_uz, source_lang, status)
                VALUES (?, ?, ?, ?, ?, ?, 'Russian', 'active')
            """, (cells[2], ru_term, cells[6], cells[2], cells[3], cells[4]))
            added += 1
        except Exception as e:
            log.debug(f"row {r_idx} skip: {e}")

    conn.commit()
    conn.close()
    log.info(f"Imported +{added} disputed_board rows")
    return {"added": added, "total_rows": len(table.rows) - 1}


def main():
    return import_docx()


if __name__ == "__main__":
    print(main())
