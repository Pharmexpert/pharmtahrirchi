"""
Import 'Изоҳли луғат' from Qonun/SP/Изоҳли луғат.docx into annotated_words table.
102 pharmacopoeia terms with uz/ru definitions + text_no reference.
"""
import os
import sqlite3
import logging
from pathlib import Path
import re

logging.basicConfig(level=logging.INFO, format="[izohli_lugat] %(message)s")
log = logging.getLogger()

DB_PATH = os.getenv("DB_PATH", os.path.join(os.path.dirname(os.path.dirname(__file__)), "pharma_editor.db"))
DATA_DIR = Path(os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "data", "izohli_lugat"
))


def extract_term(definition: str) -> str:
    """Extract the term (word before em-dash or colon)."""
    if not definition:
        return ""
    # Usually: "Термин – изоҳ" or "Термин — изоҳ" or "Термин: изоҳ"
    m = re.split(r"\s*[–—\-:]\s*", definition, maxsplit=1)
    return m[0].strip() if m else definition[:80].strip()


def main():
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed"}

    p = DATA_DIR / "izohli_lugat.docx"
    if not p.exists():
        log.warning(f"File not found: {p}")
        return {"error": "file not found"}

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    # Ensure source/frequency columns exist on annotated_words (migration)
    for col, typ in (("source", "TEXT DEFAULT ''"), ("frequency", "INTEGER DEFAULT 1"), ("text_no", "TEXT DEFAULT ''")):
        try:
            cur.execute(f"ALTER TABLE annotated_words ADD COLUMN {col} {typ}")
        except Exception:
            pass

    doc = Document(str(p))
    if not doc.tables:
        conn.close()
        return {"error": "no tables"}

    table = doc.tables[0]
    added = 0
    skipped = 0
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            continue  # header
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 4:
            continue
        text_no = cells[0]
        ru_def = cells[1]
        uz_def = cells[3]
        if not ru_def and not uz_def:
            continue

        term_ru = extract_term(ru_def)
        term_uz = extract_term(uz_def)

        try:
            # Dedupe by uz term
            cur.execute("SELECT 1 FROM annotated_words WHERE uz = ? LIMIT 1", (term_uz,))
            if cur.fetchone():
                skipped += 1
                continue
            cur.execute("""
                INSERT INTO annotated_words
                (uz, ru, description_uz, description_ru, source_lang, source, status, text_no)
                VALUES (?, ?, ?, ?, 'Uzbek', 'state_pharmacopoeia_izohli', 'active', ?)
            """, (term_uz, term_ru, uz_def, ru_def, text_no))
            added += 1
        except Exception as e:
            log.debug(f"row {r_idx}: {e}")

    conn.commit()
    conn.close()
    result = {"added": added, "skipped_duplicates": skipped, "total_rows": len(table.rows) - 1}
    log.info(result)
    return result


if __name__ == "__main__":
    print(main())
