"""
Фармакопея DOCX/PDF файлларини MD форматга конвертация.

Манбалар:
  - Давлат фармакопеяси (ўзбек, рус) — 1-3 жилд
  - Европа фармакопеяси 9.0 (PDF)

Натижа: Qonun/5/ папкасига MD файллар

Фойдаланиш:
  python scripts/convert_pharma_to_md.py
  python scripts/convert_pharma_to_md.py --source uz   # фақат ўзбекча
  python scripts/convert_pharma_to_md.py --source eu   # фақат Европа ФП
"""
import os
import sys
import re
import logging
import argparse
from pathlib import Path

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
log = logging.getLogger("pharma2md")

# Paths
BASE = Path(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
SOURCE_DIR = BASE / "Qonun" / "1. Бирламчи материал" / "Ворд"
OUTPUT_DIR = BASE / "Qonun" / "5"

# Counters
stats = {"docx_ok": 0, "docx_fail": 0, "pdf_ok": 0, "pdf_fail": 0, "skipped": 0}


def sanitize_filename(name: str) -> str:
    """Convert Cyrillic/special chars to safe filename."""
    # Remove +++ padding from filenames
    name = re.sub(r'\++', '', name)
    # Remove special chars but keep Cyrillic, Latin, digits, dots, dashes
    name = re.sub(r'[^\w\s\.\-а-яёўқғҳА-ЯЁЎҚҒҲ]', '', name, flags=re.UNICODE)
    name = name.strip()
    # Collapse whitespace
    name = re.sub(r'\s+', '_', name)
    return name[:120]  # Max length


def convert_docx_to_md(docx_path: Path, md_path: Path) -> bool:
    """Convert DOCX to Markdown using python-docx for text extraction."""
    try:
        from docx import Document
        doc = Document(str(docx_path))

        lines = []
        # Title from first heading or filename
        title = docx_path.stem
        title = re.sub(r'\++', '', title).strip()
        lines.append(f"# {title}\n")
        lines.append(f"> Манба: {docx_path.name}\n")
        lines.append("---\n")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            # Detect heading styles
            style = (para.style.name or "").lower()
            if 'heading 1' in style or 'title' in style:
                lines.append(f"\n## {text}\n")
            elif 'heading 2' in style:
                lines.append(f"\n### {text}\n")
            elif 'heading 3' in style:
                lines.append(f"\n#### {text}\n")
            elif 'list' in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        # Tables
        for table in doc.tables:
            lines.append("")
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            lines.append("")

        content = "\n".join(lines)
        if len(content.strip()) < 50:
            log.warning(f"  Бўш файл: {docx_path.name}")
            stats["skipped"] += 1
            return False

        md_path.parent.mkdir(parents=True, exist_ok=True)
        md_path.write_text(content, encoding="utf-8")
        stats["docx_ok"] += 1
        return True

    except Exception as e:
        log.error(f"  DOCX хато: {docx_path.name}: {e}")
        stats["docx_fail"] += 1
        return False


def convert_pdf_to_md(pdf_path: Path, md_path: Path) -> bool:
    """Convert PDF to Markdown using PyMuPDF (fitz) for high-quality text extraction."""
    try:
        import fitz  # PyMuPDF

        lines = []
        title = pdf_path.stem
        lines.append(f"# {title}\n")
        lines.append(f"> Манба: {pdf_path.name}\n")
        lines.append("---\n")

        doc = fitz.open(str(pdf_path))
        for i, page in enumerate(doc):
            text = page.get_text()
            if text and text.strip():
                lines.append(f"\n<!-- Page {i+1} -->\n")
                lines.append(text.strip())

            # Extract tables via pdfplumber as fallback
            try:
                import pdfplumber
                with pdfplumber.open(str(pdf_path)) as ppdf:
                    if i < len(ppdf.pages):
                        tables = ppdf.pages[i].extract_tables()
                        for table in tables:
                            if not table:
                                continue
                            lines.append("")
                            for j, row in enumerate(table):
                                cells = [(c or "").strip().replace('\n', ' ') for c in row]
                                lines.append("| " + " | ".join(cells) + " |")
                                if j == 0:
                                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
                            lines.append("")
            except Exception:
                pass  # Skip table extraction if pdfplumber fails

        doc.close()

        content = "\n".join(lines)
        if len(content.strip()) < 50:
            stats["skipped"] += 1
            return False

        md_path.parent.mkdir(parents=True, exist_ok=True)
        md_path.write_text(content, encoding="utf-8")
        stats["pdf_ok"] += 1
        return True

    except Exception as e:
        log.error(f"  PDF хато: {pdf_path.name}: {e}")
        stats["pdf_fail"] += 1
        return False


def process_directory(src_dir: Path, out_dir: Path, file_filter: str = None):
    """Recursively process all DOCX/PDF files in directory."""
    if not src_dir.exists():
        log.warning(f"Папка топилмади: {src_dir}")
        return

    for root, dirs, files in os.walk(str(src_dir)):
        # Skip temp files
        dirs[:] = [d for d in dirs if not d.startswith('~$')]

        for fname in sorted(files):
            if fname.startswith('~$'):
                continue

            fpath = Path(root) / fname
            rel = fpath.relative_to(src_dir)

            # Build output path
            safe_name = sanitize_filename(fpath.stem) + ".md"
            # Keep directory structure
            out_subdir = out_dir
            for part in rel.parent.parts:
                safe_part = sanitize_filename(part) or part
                out_subdir = out_subdir / safe_part

            md_path = out_subdir / safe_name

            if md_path.exists() and md_path.stat().st_size > 100:
                stats["skipped"] += 1
                continue

            if fname.lower().endswith('.docx'):
                if file_filter and file_filter != 'docx':
                    continue
                log.info(f"DOCX: {rel}")
                convert_docx_to_md(fpath, md_path)

            elif fname.lower().endswith('.pdf'):
                if file_filter and file_filter != 'pdf':
                    continue
                log.info(f"PDF:  {rel}")
                convert_pdf_to_md(fpath, md_path)


def main():
    parser = argparse.ArgumentParser(description="Фармакопея → MD конвертация")
    parser.add_argument("--source", choices=["uz", "ru", "eu", "all"], default="all",
                        help="uz=ДФ ўзбекча, ru=ДФ русча, eu=Европа ФП, all=барчаси")
    parser.add_argument("--type", choices=["docx", "pdf", "all"], default="all",
                        help="Фақат docx ёки pdf")
    args = parser.parse_args()

    log.info(f"Манба: {SOURCE_DIR}")
    log.info(f"Натижа: {OUTPUT_DIR}")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    file_filter = None if args.type == "all" else args.type

    if args.source in ("uz", "all"):
        log.info("\n=== ДАВЛАТ ФАРМАКОПЕЯСИ (ЎЗБЕК) ===")
        for vol in ["1-жилд/ўзбек", "2-жилд", "3-жилд"]:
            vol_dir = SOURCE_DIR / vol
            if vol_dir.exists():
                out = OUTPUT_DIR / "DF_uzb" / vol.replace("/ўзбек", "").replace("/", "_")
                process_directory(vol_dir, out, file_filter)

    if args.source in ("ru", "all"):
        log.info("\n=== ДАВЛАТ ФАРМАКОПЕЯСИ (РУС) ===")
        for vol in ["1-жилд/рус"]:
            vol_dir = SOURCE_DIR / vol
            if vol_dir.exists():
                out = OUTPUT_DIR / "DF_rus" / vol.replace("/рус", "").replace("/", "_")
                process_directory(vol_dir, out, file_filter)

    if args.source in ("eu", "all"):
        log.info("\n=== ЕВРОПА ФАРМАКОПЕЯСИ 9.0 ===")
        eu_dir = SOURCE_DIR / "Европейская Фармакопея 9.0"
        if eu_dir.exists():
            out = OUTPUT_DIR / "EP_9.0"
            process_directory(eu_dir, out, file_filter)

    log.info(f"\n=== НАТИЖАЛАР ===")
    log.info(f"DOCX муваффақият: {stats['docx_ok']}")
    log.info(f"DOCX хато:       {stats['docx_fail']}")
    log.info(f"PDF муваффақият:  {stats['pdf_ok']}")
    log.info(f"PDF хато:        {stats['pdf_fail']}")
    log.info(f"Ўтказиб юборилди: {stats['skipped']}")
    log.info(f"Жами MD файллар:  {stats['docx_ok'] + stats['pdf_ok']}")
    log.info(f"Натижа папка:     {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
