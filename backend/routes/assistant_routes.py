"""
Pharmacist Assistant — multimodal AI chat endpoint.

Features:
  - Engine selection: llama, mistral, russian, nllb, gemini, anthropic, auto
  - Tasks: chat, edit (scientific edit), translate
  - Upload: image, file (txt/pdf/docx), audio (max 50MB)
  - Language: en, ru, uz-lat, uz-cyr

Endpoints:
  POST /api/assistant/chat            — text chat with engine choice
  POST /api/assistant/edit            — scientific editing
  POST /api/assistant/translate       — multi-engine translation
  POST /api/assistant/upload          — multimodal file upload (returns extracted text + transcribe)
  GET  /api/assistant/engines         — list available engines
"""
import os
import io
from routes.rate_limit import ai_limiter, upload_limiter
import logging
from typing import Dict, Any, Optional, List

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form

from auth import get_current_user
import db

logger = logging.getLogger("assistant_routes")
router = APIRouter(prefix="/api/assistant", tags=["assistant"])

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB

# ═══════════════════════════════════════════════════
# Pharma Expert system prompt — applied to ALL engines
# ═══════════════════════════════════════════════════
PHARMA_EXPERT_PROMPT = """Сен Pharma Expert платформасининг фармацевт асистенти, фармацевтика фанлари доктори, профессор, дори воситаларини стандартлаштириш ва сифатини назорат қилиш соҳасида халқаро тажрибага эга эксперт, фармакопейщик, фармацевтика соҳаси юқори малакали мутахассисисан.

ИХТИСОСЛИГИНГ: Норматив-ҳуқуқий ҳужжатларни (фармакопея мақолаларини) халқаро стандартлар (АҚШ фармакопеяси USP, Европа фармакопеяси Ph. Eur., Ўзбекистон Республикаси Давлат фармакопеяси) асосида экспертиза қилиш.

ТИЛЛАР: Сен билан ўзбек, рус ва инглиз тилларида мурожаат қилишади. Қайси тилда савол берилса, ўша тилда жавоб бер.

УСЛУБ: Матнни илмий таржима ва илмий таҳрир қилиш билан шуғулланасан. Матндаги хатоликларни тушунчани институтнинг аълочи талабасига тушунтиргандек, механизм ва асоси билан қадамма-қадам тушунтир. Матндаги илмий ва мантиқий хатоларни бирма-бир сана. Ҳар доим холис, илмий асосланган ва педагогик маҳорат билан жавоб бер.

ВАЗИФА (умумий): Сенга юкланган ва савол тарзида берилган мураккаб фармацевтик маълумотларни ва фармакопея мақолаларни ҳам илмий тилда мисоллар билан тушунтириш, ҳамда профессионал даражада танқидий таҳлил қилиш. Ҳар қандай материал юборилса, аввал "Вазифа учун ташаккур, бажаришга киришяпман" дегин, кейин таҳлилни бошла.

ВАЗИФА (фармакопея таҳлили): Тақдим этилган фармакопея мақоласи (норматив ҳужжат) лойиҳасини чуқур илмий, техник ва эксперт назари билан таҳлил қилиб, фармакопея талабларига номувофиқликларни, мантиқий хатоликларни ва критик номувофиқликларни аниқлаш ва сифат кўрсаткичларининг илмий асосланганлигини баҳолаш.

ТАҲЛИЛ МЕЗОНЛАРИ:
- Усуллар ва спецификацияларнинг мантиқан тўғри ишлаб чиқилганлиги
- Дори воситасини тўлиқ баҳолаш учун етарли эканлиги
- Критик бўлган чинлик, миқдорий таҳлил, аралашмалар (impurities), қолдиқ эритувчилар, микробиологик тозалик нормалари ва ҳисоблашлари тўғри келтирилганлиги
- Академик даражада ичига кириб бориш

ҚЎЛЛАНИЛАДИГАН УСУЛЛАР: Фармакопея, стандартлаштириш, техник, метрологик ва асослашнинг барча усулларини қўлла.

АНИҚЛАНГАН ХАТОЛАРНИ ҚУЙИДАГИ КАТЕГОРИЯЛАР БЎЙИЧА ЖАДВАЛ КЎРИНИШИДА ТАҚДИМ ЭТ:
| № | Категория | Хатолик тавсифи | Жойи | Тавсия (тузатиш) |
|---|-----------|-----------------|------|-----------------|

КАТЕГОРИЯЛАР:
1. Критик хатолар (Critical) — хавфсизликка бевосита таҳдид солувчи
2. Жиддий хатолар (Major) — стандартларга зид, лекин тузатиш мумкин бўлган
3. Техник камчиликлар (Minor) — расмийлаштиришдаги хатолар
(Керак бўлса ўз категорияларингни ҳам қўш)

ТАҲЛИЛ ДАВОМИДА ҚУЙИДАГИЛАРГА ЭЪТИБОР ҚАРАТ:
- Таҳлил услубларининг тўғрилиги
- Усулнинг баёни: Спецификлик, чизиқлилик, қайтарилувчанлик, аниқлик (precision) ва тўғрилик (accuracy) кўрсаткичларининг етарлилиги
- Меъёрларнинг асосланганлиги: усуллар, заррачалар ўлчами, аралашмалар миқдори (impurities), миқдорий таҳлил чегараларининг хавфсизлик талабларига мослиги
- Халқаро мувофиқлик: Етакчи халқаро фармакопеялар талаблари билан уйғунлиги (гармонизация)

ЧЕКЛОВ: Оддий фойдаланувчилар платформа хавфсизлигига оид саволлар берса (масалан "паролни қандай олиш мумкин", "админ тизимга қандай киради" ва ҳоказо), буни тақиқла ва "Бу ҳақида маълумот беришим мумкин эмас. Платформа хавфсизлиги маҳфий сақланади" деб жавоб бер."""


@router.get("/engines")
async def list_engines(current_user: Dict = Depends(get_current_user)):
    """List available AI engines for user selection."""
    engines = []

    def add(key: str, label: str, mod_name: str, languages: List[str], capabilities: List[str]):
        try:
            mod = __import__(mod_name)
            avail = mod.is_available()
            mode = mod.get_mode()
            engines.append({
                "key": key,
                "label": label,
                "available": avail,
                "mode": mode,
                "model": getattr(mod, "MODEL_ID", key),
                "languages": languages,
                "capabilities": capabilities,
            })
        except Exception as e:
            engines.append({"key": key, "label": label, "available": False, "error": str(e)[:80]})

    add("llama", "Llama 3.1 8B Uzbek", "llama_engine", ["uz", "en", "ru"], ["chat", "edit", "translate"])
    add("mistral", "Mistral 7B Uzbek", "mistral_engine", ["uz", "en"], ["chat", "edit", "translate"])
    add("russian", "Sage FRED-T5 Russian", "russian_engine", ["ru"], ["edit"])
    add("nllb", "NLLB-200 Multilingual", "translator_engine", ["en", "ru", "uz"], ["translate"])
    engines.append({"key": "gemini", "label": "Google Gemini 2.0 Flash", "available": bool(os.getenv("GOOGLE_API_KEY")), "mode": "cloud", "model": "gemini-2.0-flash", "languages": ["en", "ru", "uz"], "capabilities": ["chat", "edit", "translate", "vision"]})
    engines.append({"key": "anthropic", "label": "Anthropic Claude Haiku", "available": bool(os.getenv("ANTHROPIC_API_KEY")), "mode": "cloud", "model": "claude-haiku-4-5", "languages": ["en", "ru", "uz"], "capabilities": ["chat", "edit", "translate", "vision"]})
    engines.append({"key": "auto", "label": "Авто (энг яхшиси)", "available": True, "languages": ["en", "ru", "uz"], "capabilities": ["chat", "edit", "translate"]})
    return {"engines": engines}


async def _cloud_fallback(prompt: str, system: Optional[str] = None, task: str = "chat", source_lang: str = "en", target_lang: str = "uz") -> Dict[str, Any]:
    """Guaranteed Gemini/Anthropic fallback when local engines are empty."""
    try:
        from routes.editor_routes import generate_ai_content
        if task == "edit":
            p = (system + "\n\n" if system else "") + f"Илмий-фарма муҳаррир сифатида матнни тузат:\n\n{prompt}"
        elif task == "translate":
            p = (system + "\n\n" if system else "") + f"Translate from {source_lang} to {target_lang}:\n\n{prompt}"
        else:
            p = (system + "\n\n" if system else "") + prompt
        txt = await generate_ai_content(p, prefer="cloud")
        if txt and txt.strip():
            # Detect which cloud engine actually responded
            import os as _os
            engine_name = "gemini" if _os.getenv("GOOGLE_API_KEY") else "anthropic"
            return {"text": txt.strip(), "engine": engine_name}
    except Exception as e:
        logger.warning(f"[cloud_fallback] {e}")
    return {"text": "", "engine": "none", "error": "cloud_fallback_failed"}


async def _call_engine(engine: str, prompt: str, system: Optional[str] = None, lang: str = "uz", task: str = "chat", source_lang: str = "en", target_lang: str = "uz") -> Dict[str, Any]:
    """Route through local engines; auto-fallback to cloud if local returns empty."""
    try:
        if engine == "llama":
            import llama_engine
            if llama_engine.is_available():
                try:
                    if task == "edit":
                        txt = await llama_engine.improve_text(prompt, lang)
                    elif task == "translate":
                        txt = await llama_engine.translate(prompt, source_lang, target_lang)
                    else:
                        txt = await llama_engine.generate_async(prompt, system=system, max_tokens=1024)
                    if txt and txt.strip():
                        llama_engine.learn_record(prompt, txt, kind=f"assistant_{task}")
                        return {"text": txt.strip(), "engine": "llama_" + llama_engine.get_mode()}
                except Exception as e:
                    logger.warning(f"[llama] {e}")
            # Fall through to cloud — use it transparently (don't confuse user)
            fb = await _cloud_fallback(prompt, system, task, source_lang, target_lang)
            if fb.get("text"):
                return fb
            return {"text": "Llama engine юкланмоқда (3-5 мин) ва cloud fallback ҳам ишламади.", "engine": "llama_empty", "error": "all_failed"}

        if engine == "mistral":
            import mistral_engine
            if mistral_engine.is_available():
                try:
                    if task == "edit":
                        txt = await mistral_engine.improve_text(prompt, lang)
                    elif task == "translate":
                        txt = await mistral_engine.translate(prompt, source_lang, target_lang)
                    else:
                        txt = await mistral_engine.generate_async(prompt, system=system, max_tokens=1024)
                    if txt and txt.strip():
                        return {"text": txt.strip(), "engine": "mistral_" + mistral_engine.get_mode()}
                except Exception as e:
                    logger.warning(f"[mistral] {e}")
            # Fall through to cloud — transparent
            fb = await _cloud_fallback(prompt, system, task, source_lang, target_lang)
            if fb.get("text"):
                return fb
            return {"text": "Mistral ва cloud fallback ишламади.", "engine": "mistral_empty", "error": "all_failed"}

        if engine == "russian":
            import russian_engine
            if not russian_engine.is_available():
                return {"text": "Russian engine мавжуд эмас.", "engine": "russian_unavailable", "error": "not_loaded"}
            txt = await russian_engine.improve(prompt)
            if txt and txt.strip():
                return {"text": txt.strip(), "engine": "russian_" + russian_engine.get_mode()}
            return {"text": "Russian engine бўш жавоб қайтарди.", "engine": "russian_empty", "error": "empty_response"}

        if engine == "nllb":
            import translator_engine
            if not translator_engine.is_available():
                return {"text": "NLLB engine мавжуд эмас.", "engine": "nllb_unavailable", "error": "not_loaded"}
            if task != "translate":
                return {"text": "NLLB фақат таржима учун ишлатилади. Илмий таржима режимига ўтинг.", "engine": "nllb", "error": "wrong_task"}
            txt = await translator_engine.translate_async(prompt, source_lang, target_lang)
            if txt and txt.strip():
                return {"text": txt.strip(), "engine": "nllb_" + translator_engine.get_mode()}
            return {"text": "NLLB бўш жавоб қайтарди.", "engine": "nllb_empty", "error": "empty_response"}

        if engine in ("gemini", "anthropic"):
            # Force the specific provider
            import os as _os
            if engine == "gemini" and not _os.getenv("GOOGLE_API_KEY"):
                return {"text": "Gemini API key настройкаланмаган", "engine": "gemini", "error": "no_key"}
            if engine == "anthropic" and not _os.getenv("ANTHROPIC_API_KEY"):
                return {"text": "Anthropic API key настройкаланмаган", "engine": "anthropic", "error": "no_key"}

            if task == "edit":
                p = (system + "\n\n" if system else "") + f"Қуйидаги матнни илмий-фарма муҳаррир сифатида тузат. Фақат тузатилган матнни қайтар:\n\n{prompt}"
            elif task == "translate":
                p = (system + "\n\n" if system else "") + f"Translate from {source_lang} to {target_lang}. Return only the translation:\n\n{prompt}"
            else:
                p = (system + "\n\n" if system else "") + prompt

            try:
                if engine == "gemini":
                    from routes.ai_helpers import get_gemini
                    gemini = get_gemini()
                    if gemini and hasattr(gemini, 'models'):
                        import asyncio as _aio
                        response = await _aio.get_event_loop().run_in_executor(
                            None,
                            lambda: gemini.models.generate_content(model="gemini-2.0-flash", contents=p)
                        )
                        txt = response.text if hasattr(response, 'text') else ""
                    elif gemini:
                        import asyncio as _aio
                        response = await _aio.get_event_loop().run_in_executor(None, lambda: gemini.generate_content(p))
                        txt = response.text if hasattr(response, 'text') else ""
                    else:
                        txt = ""
                else:  # anthropic
                    from routes.ai_helpers import get_anthropic
                    anthropic_client = get_anthropic()
                    if anthropic_client:
                        import asyncio as _aio
                        msg = await _aio.get_event_loop().run_in_executor(
                            None,
                            lambda: anthropic_client.messages.create(
                                model="claude-haiku-4-5-20251001",
                                max_tokens=4096,
                                messages=[{"role": "user", "content": p}]
                            )
                        )
                        txt = msg.content[0].text if msg.content else ""
                    else:
                        txt = ""

                if txt and txt.strip():
                    return {"text": txt.strip(), "engine": engine}
            except Exception as e:
                logger.exception(f"[{engine}] failed")
                return {"text": f"{engine} хато: {str(e)[:200]}", "engine": engine, "error": str(e)[:200]}
            return {"text": f"{engine} бўш жавоб", "engine": engine, "error": "empty_response"}

        if engine == "auto":
            # Try local first, then cloud as fallback
            for fallback_engine in ("llama", "mistral", "gemini", "anthropic"):
                r = await _call_engine(fallback_engine, prompt, system=system, lang=lang, task=task, source_lang=source_lang, target_lang=target_lang)
                if r.get("text") and not r.get("error"):
                    return r
            return {"text": "Барча engine'лар бўш жавоб қайтарди.", "engine": "all_empty", "error": "all_failed"}

        return {"text": f"Номаълум engine: {engine}", "engine": engine, "error": "unknown_engine"}

    except Exception as e:
        logger.warning(f"[assistant] {engine} threw: {e}")
        return {"text": f"Хато: {e}", "engine": engine, "error": str(e)}


@router.post("/history/save")
async def save_chat(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Save assistant chat history for user."""
    try:
        conn = db.connect_db()
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS assistant_chats (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT, session_id TEXT,
                title TEXT, messages TEXT, engine TEXT, lang TEXT,
                mode TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        import json as _json
        session_id = payload.get("session_id") or str(current_user.get("id", "") + "_" + str(int(__import__('time').time())))
        messages = _json.dumps(payload.get("messages", []), ensure_ascii=False)
        title = payload.get("title", (payload.get("messages") or [{}])[0].get("content", "Чат")[:60])
        # Check if session exists → update
        cur.execute("SELECT id FROM assistant_chats WHERE session_id = ? AND user_id = ?", (session_id, current_user.get("id", "")))
        row = cur.fetchone()
        if row:
            cur.execute("UPDATE assistant_chats SET messages = ?, title = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (messages, title, row[0]))
        else:
            cur.execute("""INSERT INTO assistant_chats (user_id, session_id, title, messages, engine, lang, mode) VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (current_user.get("id", ""), session_id, title, messages, payload.get("engine", ""), payload.get("lang", ""), payload.get("mode", "")))
        conn.commit()
        conn.close()
        return {"success": True, "session_id": session_id}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.get("/history/list")
async def list_chats(current_user: Dict = Depends(get_current_user)):
    """List user's saved chat sessions."""
    try:
        conn = db.connect_db()
        conn.row_factory = db.sqlite3.Row
        cur = conn.cursor()
        cur.execute("""CREATE TABLE IF NOT EXISTS assistant_chats (
            id INTEGER PRIMARY KEY AUTOINCREMENT, user_id TEXT, session_id TEXT, title TEXT,
            messages TEXT, engine TEXT, lang TEXT, mode TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )""")
        cur.execute("SELECT id, session_id, title, engine, lang, mode, created_at, updated_at FROM assistant_chats WHERE user_id = ? ORDER BY updated_at DESC LIMIT 100", (current_user.get("id", ""),))
        chats = [dict(r) for r in cur.fetchall()]
        conn.close()
        return {"chats": chats}
    except Exception as e:
        return {"chats": [], "error": str(e)}


@router.get("/history/{session_id}")
async def get_chat(session_id: str, current_user: Dict = Depends(get_current_user)):
    """Get full chat messages by session_id."""
    try:
        conn = db.connect_db()
        conn.row_factory = db.sqlite3.Row
        cur = conn.cursor()
        cur.execute("SELECT * FROM assistant_chats WHERE session_id = ? AND user_id = ?", (session_id, current_user.get("id", "")))
        row = cur.fetchone()
        conn.close()
        if not row:
            return {"found": False}
        import json as _json
        d = dict(row)
        d["messages"] = _json.loads(d.get("messages", "[]") or "[]")
        return {"found": True, "chat": d}
    except Exception as e:
        return {"found": False, "error": str(e)}


@router.delete("/history/{session_id}")
async def delete_chat(session_id: str, current_user: Dict = Depends(get_current_user)):
    """Delete a chat session."""
    try:
        conn = db.connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM assistant_chats WHERE session_id = ? AND user_id = ?", (session_id, current_user.get("id", "")))
        conn.commit()
        conn.close()
        return {"success": True}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/chat")
async def chat(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """General-purpose chat with engine selection. Rate limit: 20/min per user."""
    user_key = str(current_user.get("id", "anon"))
    if not ai_limiter.is_allowed(user_key):
        raise HTTPException(status_code=429, detail="Сўровлар лимити ошди (20/мин). 1 минутдан сўнг қайта уриниб кўринг.")
    message = (payload.get("message") or "").strip()
    if not message:
        return {"text": ""}
    engine = payload.get("engine", "auto")
    lang = payload.get("lang", "uz")
    history = payload.get("history", [])  # [{role, content}]
    # ALWAYS use Pharma Expert prompt as the base; user can append extra instructions via `system`
    extra_system = payload.get("system", "")
    system = PHARMA_EXPERT_PROMPT + (("\n\nҚўшимча йўриқнома: " + extra_system) if extra_system else "")

    # For Russian engine — don't include history (it corrects plain text, not dialog)
    # For NLLB — wrong mode, return error
    if engine == "nllb":
        return {"text": "NLLB фақат таржима учун ишлатилади. Илмий таржима режимига ўтинг.", "engine": "nllb", "error": "wrong_task"}
    if engine == "russian":
        # Only pass current message, no history, no system prompt (it's a T5 corrector not a chat)
        res = await _call_engine(engine, message, lang=lang, task="chat")
        return res

    # For LLMs with dialog capability — build proper chat context
    context = ""
    if history:
        for h in history[-6:]:  # last 3 turns
            role = h.get("role", "user")
            content = h.get("content", "")
            if content and len(content) < 1000:  # skip huge attachments
                context += f"{'Фойдаланувчи' if role == 'user' else 'Ёрдамчи'}: {content}\n\n"
    full_message = (context + f"Фойдаланувчи: {message}\n\nЁрдамчи:") if context else message

    res = await _call_engine(engine, full_message, system=system, lang=lang, task="chat")
    return res


@router.post("/edit")
async def edit(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Scientific editing — Pharma Expert prompt applied. Rate: 20/min."""
    user_key = str(current_user.get("id", "anon"))
    if not ai_limiter.is_allowed(user_key):
        raise HTTPException(status_code=429, detail="Сўровлар лимити ошди (20/мин).")
    text = (payload.get("text") or "").strip()
    engine = payload.get("engine", "auto")
    lang = payload.get("lang", "uz")
    if not text:
        return {"text": ""}
    # TM check
    try:
        import tm_search
        tm = tm_search.best_match(text, src_lang=f"{lang}_raw", tgt_lang=f"{lang}_improved")
        if tm and tm["score"] >= 0.95:
            return {"text": tm["tgt_text"], "engine": "tm_exact"}
    except Exception:
        pass
    result = await _call_engine(engine, text, system=PHARMA_EXPERT_PROMPT, lang=lang, task="edit")
    # TM learn
    out_text = result.get("text", "")
    if out_text and out_text != text:
        try:
            import tm_search
            tm_search.save_to_tm(text, out_text, src_lang=f"{lang}_raw", tgt_lang=f"{lang}_improved", source="assistant_edit")
        except Exception:
            pass
    return result


@router.post("/translate")
async def translate(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Translation — Pharma Expert prompt applied. Rate: 20/min."""
    user_key = str(current_user.get("id", "anon"))
    if not ai_limiter.is_allowed(user_key):
        raise HTTPException(status_code=429, detail="Сўровлар лимити ошди (20/мин).")
    text = (payload.get("text") or "").strip()
    engine = payload.get("engine", "auto")
    src = payload.get("source_lang", "en")
    tgt = payload.get("target_lang", "uz")
    if not text:
        return {"text": ""}
    base_src = "ru" if src == "ru" else ("uz" if src.startswith("uz") else "en")
    base_tgt = "ru" if tgt == "ru" else ("uz" if tgt.startswith("uz") else "en")
    # TM check
    try:
        import tm_search
        tm = tm_search.best_match(text, src_lang=base_src, tgt_lang=base_tgt)
        if tm and tm["score"] >= 0.95:
            return {"text": tm["tgt_text"], "engine": "tm_exact"}
    except Exception:
        pass
    result = await _call_engine(engine, text, system=PHARMA_EXPERT_PROMPT, source_lang=src, target_lang=tgt, task="translate")
    # TM learn
    out_text = result.get("text", "")
    if out_text and out_text != text:
        try:
            import tm_search
            tm_search.save_to_tm(text, out_text, src_lang=base_src, tgt_lang=base_tgt, source="assistant_translate")
        except Exception:
            pass
    return result


@router.post("/upload")
async def upload(file: UploadFile = File(...), kind: str = Form("auto"), current_user: Dict = Depends(get_current_user)):
    # Upload limiter: 5 files/minute
    user_key = str(current_user.get("id", "anon"))
    if not upload_limiter.is_allowed(user_key):
        raise HTTPException(status_code=429, detail="Файл юклаш лимити ошди (5/мин).")
    """
    Multimodal upload (image, file, audio).
    Returns extracted text + metadata.

    Limits:
      - Max 50 MB per file
      - Allowed: txt, md, csv, pdf, docx, jpg, png, webp, gif, mp3, wav, m4a, ogg
    """
    name = (file.filename or "").lower()
    raw = await file.read()
    if len(raw) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"Файл ўлчами 50 МБдан катта ({len(raw) // (1024*1024)} МБ)")
    if not raw:
        raise HTTPException(status_code=400, detail="Бўш файл")

    info = {"filename": name, "size": len(raw), "size_mb": round(len(raw) / (1024 * 1024), 2)}

    # Text files
    if name.endswith(('.txt', '.md', '.csv')):
        for enc in ('utf-8', 'utf-8-sig', 'cp1251', 'latin-1'):
            try:
                info["text"] = raw.decode(enc)
                info["kind"] = "text"
                return info
            except UnicodeDecodeError:
                continue
        info["text"] = raw.decode('utf-8', errors='ignore')
        info["kind"] = "text"
        return info

    # DOCX
    if name.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            info["text"] = "\n".join(paragraphs)
            info["kind"] = "docx"
            return info
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX ўқишда хато: {e}")

    # PDF
    if name.endswith('.pdf'):
        text = ""
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                text = "\n".join((p.extract_text() or "") for p in pdf.pages)
        except Exception:
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(raw))
                text = "\n".join((p.extract_text() or "") for p in reader.pages)
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"PDF ўқишда хато: {e}")
        info["text"] = text
        info["kind"] = "pdf"
        return info

    # Image — return base64 + dimensions, can be sent to vision-capable engine
    if name.endswith(('.jpg', '.jpeg', '.png', '.webp', '.gif')):
        import base64
        info["base64"] = base64.b64encode(raw).decode('ascii')
        info["mime"] = "image/jpeg" if name.endswith(('.jpg', '.jpeg')) else f"image/{name.rsplit('.', 1)[-1]}"
        info["kind"] = "image"
        info["text"] = f"[Image: {name}, {info['size_mb']} MB]"
        return info

    # Audio — return base64 + try transcription if Whisper available
    if name.endswith(('.mp3', '.wav', '.m4a', '.ogg', '.flac')):
        import base64
        info["base64"] = base64.b64encode(raw).decode('ascii')
        info["mime"] = f"audio/{name.rsplit('.', 1)[-1]}"
        info["kind"] = "audio"
        # Try Whisper transcription
        try:
            import whisper_engine
            if whisper_engine.is_available():
                transcript = await whisper_engine.transcribe_async(raw, name)
                info["text"] = transcript or f"[Audio: {name}]"
                info["transcribed"] = bool(transcript)
            else:
                info["text"] = f"[Audio: {name}, {info['size_mb']} MB — transcription not available]"
        except Exception:
            info["text"] = f"[Audio: {name}, {info['size_mb']} MB]"
        return info

    raise HTTPException(status_code=400, detail=f"Қўллаб-қувватланмаган формат: {name}")
