"""
Tilshunos (Linguist) integrated workflow endpoints.

This is the unified text-improvement pipeline that combines all platform tools:
  1. Translation (EN ↔ RU ↔ UZ via AI)
  2. Spell-check (Hunspell + tahrirchi.db)
  3. Morphology (morphology.py)
  4. Grammar checker (grammar_checker.py)
  5. Sayqallash rules (sayqallash_rules)
  6. Pattern-based linguistic rules (uzbek_linguistic_rules)
  7. Self-learning loop (db.learn_from_correction)

User submits 4 text panels (3 languages + outcome) → backend returns
unified analysis with categorized issues, suggestions, and confidence scores.
"""
import logging
import os
import re
from typing import Dict, Any, List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File


def _detect_text_script(text: str) -> str:
    """Detect if text is primarily Cyrillic or Latin."""
    cyr = sum(1 for c in text[:200] if '\u0400' <= c <= '\u04FF')
    lat = sum(1 for c in text[:200] if 'a' <= c.lower() <= 'z')
    return "cyr" if cyr >= lat else "lat"


def _to_text_script(value: str, target_script: str) -> str:
    """Convert value to match target script."""
    if not value or not value.strip():
        return value
    val_cyr = any('\u0400' <= c <= '\u04FF' for c in value)
    val_lat = any('a' <= c.lower() <= 'z' for c in value)
    if target_script == "cyr" and val_lat and not val_cyr:
        try:
            import dual_script
            return dual_script.to_cyrillic(value) or value
        except Exception:
            pass
    elif target_script == "lat" and val_cyr and not val_lat:
        try:
            import dual_script
            return dual_script.to_latin(value) or value
        except Exception:
            pass
    return value

from auth import get_current_user
import db

logger = logging.getLogger("tilshunos_routes")
router = APIRouter(prefix="/api/tilshunos", tags=["tilshunos"])


@router.post("/extract-text")
async def extract_text_from_file(file: UploadFile = File(...), current_user: Dict = Depends(get_current_user)):
    """Extract plain text from uploaded .txt/.docx/.pdf for Tilshunos editor."""
    name = (file.filename or "").lower()
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Empty file")
    try:
        if name.endswith(".txt") or name.endswith(".md") or name.endswith(".csv"):
            for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
                try:
                    return {"text": raw.decode(enc)}
                except UnicodeDecodeError:
                    continue
            return {"text": raw.decode("utf-8", errors="ignore")}
        if name.endswith(".docx"):
            import io
            from docx import Document
            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            return {"text": "\n".join(paragraphs)}
        if name.endswith(".pdf"):
            import io
            text = ""
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(raw)) as pdf:
                    text = "\n".join((p.extract_text() or "") for p in pdf.pages)
            except Exception:
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(io.BytesIO(raw))
                    text = "\n".join((p.extract_text() or "") for p in reader.pages)
                except Exception:
                    from PyPDF2 import PdfReader  # type: ignore
                    reader = PdfReader(io.BytesIO(raw))
                    text = "\n".join((p.extract_text() or "") for p in reader.pages)
            return {"text": text}
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {name}")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("extract-text failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/rules/stats")
async def get_rules_stats(current_user: Dict = Depends(get_current_user)):
    """
    Return total counts of all linguistic rules across the platform.

    Used by dashboard to show: orthography rules, punctuation rules, etc.
    """
    try:
        from uzbek_linguistic_rules import get_rules_summary
        canonical_summary = get_rules_summary()
    except Exception as e:
        canonical_summary = {"error": str(e)}

    # Sayqallash rules from DB
    conn = db.connect_db()
    cur = conn.cursor()
    sayqallash_total = 0
    sayqallash_by_type = {}
    try:
        cur.execute("SELECT error_type, COUNT(*) FROM sayqallash_rules GROUP BY error_type")
        for row in cur.fetchall():
            sayqallash_by_type[row[0] or "unknown"] = row[1]
        sayqallash_total = sum(sayqallash_by_type.values())
    except Exception:
        pass

    # Hunspell stats
    hunspell_stats = {}
    try:
        import os
        HUNSPELL_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "hunspell")
        for fname, key in [("uz_UZ.dic", "lat_words"), ("uz_UZ_Cyrl.dic", "cyr_words"),
                            ("uz_UZ.aff", "lat_aff"), ("uz_UZ_Cyrl.aff", "cyr_aff")]:
            path = os.path.join(HUNSPELL_DIR, fname)
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
                if fname.endswith(".dic"):
                    hunspell_stats[key] = len([l for l in content.split("\n") if l.strip() and not l.strip().isdigit()])
                else:
                    hunspell_stats[key + "_sfx"] = content.count("SFX ")
                    hunspell_stats[key + "_pfx"] = content.count("PFX ")
                    hunspell_stats[key + "_rep"] = content.count("REP ")
    except Exception:
        pass

    # Tahrirchi lexicon
    tahrirchi_words = 0
    try:
        import sqlite3
        if hasattr(db, "TAHRIRCHI_DB_PATH"):
            import os as _os
            if _os.path.exists(db.TAHRIRCHI_DB_PATH):
                tc = sqlite3.connect(db.TAHRIRCHI_DB_PATH)
                tahrirchi_words = tc.execute("SELECT COUNT(*) FROM dictionary").fetchone()[0]
                tc.close()
    except Exception:
        pass

    # Counts from other tables
    counts = {}
    for table in ("synonyms", "annotated_words", "disputed_words", "abbreviations", "alignments"):
        try:
            counts[table] = cur.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0]
        except Exception:
            counts[table] = 0
    conn.close()

    return {
        "canonical_rules": canonical_summary,
        "sayqallash_total": sayqallash_total,
        "sayqallash_by_type": sayqallash_by_type,
        "hunspell": hunspell_stats,
        "tahrirchi_lexicon_words": tahrirchi_words,
        "platform_databases": counts,
    }


@router.post("/check")
async def comprehensive_check(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """
    Comprehensive text check across all 5 dimensions.

    Request:
        {
            "text": "...",
            "lang": "uz" | "ru" | "en",
            "categories": ["orthography", "punctuation", "syntax", "morphology", "grammar"]
                          (optional, defaults to all)
        }

    Response:
        {
            "text": "...",
            "lang": "uz",
            "issues": [
                {
                    "category": "orthography|punctuation|syntax|morphology|grammar",
                    "subcategory": "...",
                    "rule_id": "ORTH-001" | "GRAM-005" | "RULE_DB:1234" | "MORPH",
                    "error_type": "S/Apostrophe",
                    "from_index": N, "to_index": M,
                    "matched_text": "...",
                    "suggestion": "...",
                    "message": "...",
                    "severity": "low|medium|high",
                    "source": "pattern|sayqallash|morphology|grammar_checker"
                }
            ],
            "by_category": { "orthography": N, ... },
            "by_severity": { "high": X, "medium": Y, "low": Z },
            "total": N,
            "morphology": [...],
            "score": 0.0-1.0  (text quality estimate)
        }
    """
    text = payload.get("text", "").strip()
    lang = payload.get("lang", "uz")
    requested_categories = payload.get("categories", None)

    if not text:
        raise HTTPException(status_code=400, detail="text is required")
    if len(text) > 50000:
        raise HTTPException(status_code=400, detail="text too long (max 50000)")

    all_issues = []
    # Track which word spans already have issues to avoid duplicates
    seen_spans = set()  # (from_index, to_index)

    # Uzbek-aware tokenizer: include straight apostrophe (') curly (\u2019)
    # and ʻ (\u02bb) and ʼ (\u02bc) as part of words so "o'quvchilaga" stays whole.
    UZ_WORD_RE = r"[\w'\u2019\u02bb\u02bc\u2018]+"

    # ───── 1. Per-word Hunspell spell check (Uzbek) ─────
    if lang == "uz":
        try:
            import spellcheck
            # Detect script by checking any cyrillic character
            is_cyr = any("\u0400" <= ch <= "\u04FF" for ch in text)

            # Cross-script transliteration helper — checks the SAME word in BOTH scripts
            # so a word marked unknown in Cyrillic Hunspell but known in Latin (or vice versa)
            # is treated as correct. Invisible to user.
            def _translit_to_other(w: str, from_cyr: bool) -> str:
                try:
                    import transliterate as tl  # local helper module
                    if from_cyr:
                        return tl.to_latin(w)
                    return tl.to_cyrillic(w)
                except Exception:
                    return w

            for match in re.finditer(UZ_WORD_RE, text, re.UNICODE):
                word = match.group()
                # Strip surrounding apostrophes (not part of word) but keep inner ones
                word = word.strip("'\u2019\u02bb\u02bc\u2018")
                if len(word) < 2:
                    continue
                # Skip digits / mixed
                if any(ch.isdigit() for ch in word):
                    continue

                try:
                    # Uzbek Hunspell dict uses ʻ (\u02bb modifier letter). Users often
                    # type ' (straight) or ' (\u2019 curly). Try ALL apostrophe variants.
                    variants = [word]
                    normalized_straight = word.replace("\u2019", "'").replace("\u02bb", "'").replace("\u02bc", "'").replace("\u2018", "'")
                    normalized_hunspell = word.replace("'", "\u02bb").replace("\u2019", "\u02bb").replace("\u02bc", "\u02bb").replace("\u2018", "\u02bb")
                    if normalized_straight != word:
                        variants.append(normalized_straight)
                    if normalized_hunspell != word:
                        variants.append(normalized_hunspell)

                    ok = any(spellcheck.is_correct(v, is_cyrillic=is_cyr) for v in variants)
                    if ok:
                        continue

                    # Cross-script check: try the SAME word in the OTHER script's dictionary
                    other_word = _translit_to_other(word, is_cyr)
                    if other_word and other_word != word:
                        if spellcheck.is_correct(other_word, is_cyrillic=not is_cyr):
                            continue
                        # also try apostrophe-normalized for the other script
                        ow_h = other_word.replace("'", "\u02bb").replace("\u2019", "\u02bb")
                        if ow_h != other_word and spellcheck.is_correct(ow_h, is_cyrillic=not is_cyr):
                            continue

                    # Try suggestions from all variants
                    suggestions = []
                    for v in variants:
                        s = spellcheck.suggest(v, is_cyrillic=is_cyr, max_suggestions=5)
                        suggestions.extend(s)
                    # Dedupe while preserving order, convert all to straight apostrophe for UX
                    seen = set()
                    clean_suggestions = []
                    for s in suggestions:
                        s_clean = s.replace("\u02bb", "'").replace("\u2019", "'").replace("\u02bc", "'").replace("\u2018", "'")
                        if s_clean not in seen and s_clean != normalized_straight:
                            seen.add(s_clean)
                            clean_suggestions.append(s_clean)
                    suggestions = clean_suggestions[:5]
                    if not suggestions:
                        continue

                    span = (match.start(), match.end())
                    if span in seen_spans:
                        continue
                    seen_spans.add(span)

                    all_issues.append({
                        "category": "orthography",
                        "subcategory": "spelling",
                        "rule_id": "HUNSPELL",
                        "error_type": "S/Spelling",
                        "from_index": match.start(),
                        "to_index": match.end(),
                        "matched_text": word,
                        "suggestion": suggestions[0] if suggestions else "",
                        "suggestions": suggestions,
                        "message": f"Имло хатоси: '{word}' луғатда топилмади",
                        "severity": "high",
                        "source": "hunspell",
                    })
                except Exception:
                    continue
        except Exception as e:
            logger.warning(f"[tilshunos] hunspell check failed: {e}")

    # ───── 2. Sayqallash rules (DB-based exact + semantic) ─────
    if lang == "uz" and (not requested_categories or "orthography" in requested_categories or "spelling" in requested_categories):
        try:
            sayqallash_results = db.get_rules_for_text(text, lang) or []
            for r in sayqallash_results:
                span = (r.get("from_index", 0), r.get("to_index", 0))
                if span in seen_spans:
                    continue
                seen_spans.add(span)

                all_issues.append({
                    "category": "orthography",
                    "subcategory": "sayqallash_db",
                    "rule_id": f"RULE_DB:{r.get('frequency', 0)}",
                    "error_type": r.get("error_type", "S/Spelling"),
                    "from_index": r.get("from_index", 0),
                    "to_index": r.get("to_index", 0),
                    "matched_text": r.get("old_value", ""),
                    "suggestion": r.get("new_value", ""),
                    "suggestions": [r.get("new_value", "")],
                    "message": f"Sayqallash қоидаси: {r.get('error_type', '')}",
                    "severity": "medium",
                    "source": "sayqallash_db",
                })
        except Exception as e:
            logger.warning(f"[tilshunos] sayqallash check failed: {e}")

    # ───── 3. Sentence-level pattern rules (orthography/punctuation sentence-wide only) ─────
    try:
        from uzbek_linguistic_rules import check_text_against_rules
        # Only sentence-level rules, not single-char patterns that cause false positives
        violations = check_text_against_rules(
            text,
            categories=["punctuation"],  # only punctuation patterns are reliable
            script="both",
        )
        for v in violations:
            span = (v.from_index, v.to_index)
            if span in seen_spans:
                continue
            seen_spans.add(span)
            issue = v.to_dict()
            issue["source"] = "pattern_rules"
            issue["suggestions"] = [v.suggestion] if v.suggestion else []
            all_issues.append(issue)
    except Exception as e:
        logger.warning(f"[tilshunos] pattern rules check failed: {e}")

    # ───── 3. Morphology analysis ─────
    morphology_results = []
    if lang == "uz" and (not requested_categories or "morphology" in requested_categories):
        try:
            import morphology
            analyzer = morphology.get_analyzer()
            for word in re.findall(r"\w+", text, re.UNICODE):
                if len(word) < 2:
                    continue
                analysis = analyzer.analyze(word)
                if analysis:
                    morphology_results.append(analysis.to_dict())
                    if not analysis.valid_order:
                        all_issues.append({
                            "category": "morphology",
                            "subcategory": "slot_order",
                            "rule_id": "MORPH-ORDER",
                            "error_type": "G/Morphology",
                            "from_index": text.find(word),
                            "to_index": text.find(word) + len(word),
                            "matched_text": word,
                            "suggestion": "",
                            "message": " | ".join(analysis.order_issues),
                            "severity": "medium",
                            "source": "morphology",
                        })
        except Exception as e:
            logger.warning(f"[tilshunos] morphology failed: {e}")

    # ───── 4. Grammar checker ─────
    if lang == "uz" and (not requested_categories or "grammar" in requested_categories):
        try:
            import grammar_checker
            checker = grammar_checker.get_checker()
            grammar_issues = checker.check(text, lang=lang)
            for gi in grammar_issues:
                all_issues.append({
                    "category": "grammar",
                    "subcategory": gi.issue_type.split("/")[-1] if "/" in gi.issue_type else gi.issue_type,
                    "rule_id": f"GRAM-CHECKER",
                    "error_type": gi.issue_type,
                    "from_index": gi.from_index,
                    "to_index": gi.to_index,
                    "matched_text": gi.word,
                    "suggestion": ", ".join(gi.suggestions[:3]) if gi.suggestions else "",
                    "message": gi.message,
                    "severity": "high" if gi.confidence > 0.7 else "medium",
                    "source": "grammar_checker",
                })
        except Exception as e:
            logger.warning(f"[tilshunos] grammar check failed: {e}")

    # Ensure all suggestions match input text's script (Cyr↔Lat)
    text_script = _detect_text_script(text)
    for issue in all_issues:
        if issue.get("suggestion"):
            parts = [_to_text_script(s.strip(), text_script) for s in issue["suggestion"].split(",")]
            issue["suggestion"] = ", ".join(parts)
        if issue.get("suggestions"):
            issue["suggestions"] = [_to_text_script(s, text_script) for s in issue["suggestions"]]

    # Aggregate stats
    by_category = {}
    by_severity = {"low": 0, "medium": 0, "high": 0}
    for issue in all_issues:
        cat = issue.get("category", "unknown")
        by_category[cat] = by_category.get(cat, 0) + 1
        sev = issue.get("severity", "medium")
        by_severity[sev] = by_severity.get(sev, 0) + 1

    # Quality score: 1.0 if no issues, decrease with high-severity issues
    word_count = len(text.split())
    if word_count == 0:
        score = 1.0
    else:
        weight = by_severity["high"] * 3 + by_severity["medium"] * 1 + by_severity["low"] * 0.3
        score = max(0.0, 1.0 - (weight / max(1, word_count)))

    return {
        "text": text,
        "lang": lang,
        "issues": all_issues,
        "by_category": by_category,
        "by_severity": by_severity,
        "total": len(all_issues),
        "morphology": morphology_results[:30],  # cap to avoid huge response
        "score": round(score, 3),
        "word_count": word_count,
    }


@router.post("/classify")
async def classify_words(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """
    Classify every word in text by its source category.

    Returns spans with one of:
      - "unknown"    → not in any dict (red error)
      - "annotated"  → in annotated_words table (teal #0891B2)
      - "disputed"   → in disputed_words table (orange #F97316)
      - "abbreviation" → in abbreviations table (indigo #6366F1)
      - "known"      → in Hunspell dict (dark blue #1E40AF)
    """
    text = payload.get("text", "")
    lang = payload.get("lang", "uz")

    if not text:
        return {"spans": [], "counts": {}}

    # Load all "known-in-platform" word sets once
    annotated_set = set()
    disputed_set = set()
    abbr_set = set()
    try:
        conn = db.connect_db()
        cur = conn.cursor()
        for table, target in (("annotated_words", annotated_set),
                              ("disputed_words", disputed_set),
                              ("abbreviations", abbr_set)):
            try:
                cur.execute(f"SELECT * FROM {table}")
                cols = [d[0] for d in cur.description]
                for row in cur.fetchall():
                    row_dict = dict(zip(cols, row))
                    for key in ("uz", "ru", "en", "word", "short_form", "abbreviation"):
                        val = row_dict.get(key)
                        if val and isinstance(val, str):
                            target.add(val.strip().lower())
            except Exception:
                pass
        conn.close()
    except Exception as e:
        logger.warning(f"[classify] db load failed: {e}")

    # Tokenize
    UZ_WORD_RE = r"[\w'\u2019\u02bb\u02bc\u2018]+"
    import spellcheck

    is_cyr = any("\u0400" <= ch <= "\u04FF" for ch in text)
    spans = []
    counts = {"known": 0, "annotated": 0, "disputed": 0, "abbreviation": 0, "unknown": 0}

    for match in re.finditer(UZ_WORD_RE, text, re.UNICODE):
        word_raw = match.group()
        word = word_raw.strip("'\u2019\u02bb\u02bc\u2018")
        if len(word) < 2:
            continue
        if any(ch.isdigit() for ch in word):
            continue
        wl = word.lower()
        category = "unknown"
        if wl in annotated_set:
            category = "annotated"
        elif wl in disputed_set:
            category = "disputed"
        elif wl in abbr_set:
            category = "abbreviation"
        else:
            # Check Hunspell in CURRENT script
            variants = [word, word.replace("'", "\u02bb"), word.replace("\u2019", "\u02bb")]
            if any(spellcheck.is_correct(v, is_cyrillic=is_cyr) for v in variants):
                category = "known"
            else:
                # Cross-script: Uzbek Cyrillic ↔ Latin are the same language, only different scripts.
                # Transliterate the word and check Hunspell in the OTHER script.
                try:
                    import transliterate as tl
                    other = tl.to_latin(word) if is_cyr else tl.to_cyrillic(word)
                    if other and other != word:
                        ov = [other, other.replace("'", "\u02bb"), other.replace("\u2019", "\u02bb")]
                        if any(spellcheck.is_correct(v, is_cyrillic=not is_cyr) for v in ov):
                            category = "known"
                except Exception:
                    pass
        spans.append({
            "from": match.start(),
            "to": match.end(),
            "word": word_raw,
            "category": category,
        })
        counts[category] = counts.get(category, 0) + 1

    return {"spans": spans, "counts": counts}


@router.post("/improve")
async def improve_text(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """
    Apply auto-fixes for high-confidence issues and return improved text.

    Request: { "text": "...", "lang": "uz", "auto_apply_severity": "high" }
    Response: { "improved_text": "...", "applied_fixes": [...], "remaining_issues": [...] }
    """
    text = payload.get("text", "").strip()
    lang = payload.get("lang", "uz")
    auto_severity = payload.get("auto_apply_severity", "high")

    if not text:
        raise HTTPException(status_code=400, detail="text is required")

    # First check
    check_payload = {"text": text, "lang": lang}
    check_result = await comprehensive_check(check_payload, current_user)

    issues = check_result["issues"]
    severity_threshold = {"low": 0, "medium": 1, "high": 2}
    threshold = severity_threshold.get(auto_severity, 2)

    # Sort by from_index DESC so we replace from end to avoid index shifts
    fixable = [
        i for i in issues
        if i.get("suggestion") and severity_threshold.get(i.get("severity", "medium"), 1) >= threshold
    ]
    fixable.sort(key=lambda x: -x["from_index"])

    improved = text
    applied = []
    for issue in fixable:
        start = issue["from_index"]
        end = issue["to_index"]
        old = issue["matched_text"]
        new = issue["suggestion"]
        if not new or new == old:
            continue
        improved = improved[:start] + new + improved[end:]
        applied.append({
            "from": start, "to": end, "old": old, "new": new,
            "rule_id": issue.get("rule_id"), "category": issue.get("category"),
        })

    return {
        "original_text": text,
        "improved_text": improved,
        "applied_fixes": applied,
        "fix_count": len(applied),
        "remaining_issues": [i for i in issues if i not in fixable],
    }


@router.post("/confirm")
async def confirm_correction(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """
    User confirms a correction → triggers self-learning loop.

    Request: { "wrong": "...", "correct": "...", "context": "...", "category": "...", "lang": "uz" }
    Response: result of db.learn_from_correction
    """
    wrong = payload.get("wrong", "").strip()
    correct = payload.get("correct", "").strip()
    if not wrong or not correct:
        raise HTTPException(status_code=400, detail="wrong and correct are required")

    return db.learn_from_correction(
        wrong=wrong,
        correct=correct,
        user_id=current_user.get("id", ""),
        user_name=current_user.get("name", ""),
        context=payload.get("context", ""),
        error_type=payload.get("category", "S/Spelling"),
        lang=payload.get("lang", "uz"),
        source="tilshunos_confirmed",
    )


@router.post("/translate")
async def translate_text(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """AI-translate text. Routing:
       - en/uz pair  → Mistral (best for Uzbek)
       - any pair with Russian → NLLB-200 (best for ru ↔ {en,uz})
       - fallback → Mistral or Gemini/Claude
    """
    text = (payload.get("text") or "").strip()
    src = payload.get("source_lang", "en")
    tgt = payload.get("target_lang", "uz")
    if not text:
        return {"translated": ""}

    base_src = "ru" if src == "ru" else ("uz" if src.startswith("uz") else "en")
    base_tgt = "ru" if tgt == "ru" else ("uz" if tgt.startswith("uz") else "en")
    needs_russian = "ru" in (base_src, base_tgt)

    try:
        # ──── TM (Translation Memory) check ────
        try:
            import tm_search
            tm_match = tm_search.best_match(text, src_lang=base_src, tgt_lang=base_tgt)
            if tm_match and tm_match["score"] >= 0.80:
                logger.info(f"[translate] TM exact hit ({tm_match['score']}) — skipping AI")
                return {"translated": tm_match["tgt_text"], "engine": "tm_exact", "tm_score": tm_match["score"]}
        except Exception as e:
            logger.debug(f"[translate] TM search skip: {e}")
        # ──── End TM ────

        # 0. Tahrirchi dilmash — native Uzbek model (best for uz/en/ru/kaa pairs)
        try:
            import tahrirchi_engine
            if tahrirchi_engine.is_available():
                tr = await tahrirchi_engine.translate_async(text, src, tgt)
                if tr and len(tr.strip()) > 1 and tr.strip() != text.strip():
                    return {"translated": tr.strip(), "engine": "tahrirchi_" + tahrirchi_engine.get_mode()}
        except Exception as e:
            logger.warning(f"[translate] Tahrirchi failed: {e}")

        # 1. NLLB-200 — for any pair involving Russian
        if needs_russian:
            try:
                import translator_engine
                if translator_engine.is_available() and translator_engine.supports(src, tgt):
                    tr = await translator_engine.translate_async(text, src, tgt)
                    if tr and len(tr.strip()) > 1:
                        return {"translated": tr.strip(), "engine": "nllb_" + translator_engine.get_mode()}
            except Exception as e:
                logger.warning(f"[translate] NLLB failed: {e}")

        # 2. Llama 3.1 — preferred for non-Russian pairs (better instruction following)
        try:
            import llama_engine
            if llama_engine.is_available():
                tr = await llama_engine.translate(text, src, tgt)
                if tr and len(tr.strip()) > 1:
                    llama_engine.learn_record(text, tr, kind="llama_translate")
                    return {"translated": tr.strip(), "engine": "llama_" + llama_engine.get_mode()}
        except Exception as e:
            logger.warning(f"[translate] Llama failed: {e}")

        # 3. Mistral — fallback for Uzbek and other pairs
        try:
            import mistral_engine
            if mistral_engine.is_available():
                tr = await mistral_engine.translate(text, src, tgt)
                if tr and len(tr.strip()) > 1:
                    return {"translated": tr.strip(), "engine": "mistral_" + mistral_engine.get_mode()}
        except Exception as e:
            logger.warning(f"[translate] Mistral path failed: {e}")
        # Cloud fallback
        from routes.editor_routes import generate_ai_content
        label = {"en": "English", "ru": "Russian", "uz": "Uzbek (Latin)", "uz-lat": "Uzbek (Latin)", "uz-cyr": "Uzbek (Cyrillic)"}
        prompt = f"Translate the following {label.get(src, src)} text to {label.get(tgt, tgt)}. Return ONLY the translation, no explanations.\n\n{text}"
        translated = await generate_ai_content(prompt, prefer="cloud")
        result_text = (translated or "").strip()
        # ──── TM learn: save successful AI translation for future reuse ────
        if result_text and len(result_text) > 2:
            try:
                import tm_search
                tm_search.save_to_tm(text, result_text, src_lang=base_src, tgt_lang=base_tgt, source="ai_cloud")
            except Exception:
                pass
        return {"translated": result_text, "engine": "cloud"}
    except Exception as e:
        logger.exception("translate failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/ai-improve")
async def ai_improve(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """AI scientific edit. Routing by lang:
       - ru → sage-fredt5-large (Russian-specific corrector)
       - uz → Mistral-7B-Instruct-Uz
       - en/other → Mistral / cloud
    """
    text = (payload.get("text") or "").strip()
    lang = payload.get("lang", "uz")
    if not text:
        return {"improved": ""}
    try:
        # ──── TM check: if we've seen this exact text improved before, reuse it ────
        try:
            import tm_search
            # For editing, we search TM with src_lang=lang, tgt_lang=lang (same-lang improvement)
            tm_match = tm_search.best_match(text, src_lang=f"{lang}_raw", tgt_lang=f"{lang}_improved")
            if tm_match and tm_match["score"] >= 0.80:
                logger.info(f"[ai-improve] TM exact hit ({tm_match['score']}) — skipping AI")
                return {"improved": tm_match["tgt_text"], "engine": "tm_exact", "tm_score": tm_match["score"]}
        except Exception as e:
            logger.debug(f"[ai-improve] TM search skip: {e}")
        # ──── End TM ────

        # Russian → use specialized Russian corrector
        if lang == "ru" or lang.startswith("ru"):
            try:
                import russian_engine
                if russian_engine.is_available():
                    improved = await russian_engine.improve(text)
                    if improved and improved != text:
                        return {"improved": improved, "engine": "russian_" + russian_engine.get_mode()}
            except Exception as e:
                logger.warning(f"[ai-improve] Russian engine failed: {e}")

        # Try Llama 3.1 first (better than Mistral 7B)
        try:
            import llama_engine
            if llama_engine.is_available():
                improved = await llama_engine.improve_text(text, lang)
                if improved and improved != text:
                    llama_engine.learn_record(text, improved, kind="llama_improve")
                    return {"improved": improved, "engine": "llama_" + llama_engine.get_mode()}
        except Exception as e:
            logger.warning(f"[ai-improve] Llama engine failed: {e}")

        import mistral_engine
        if mistral_engine.is_available():
            improved = await mistral_engine.improve_text(text, lang)
            return {"improved": improved or text, "engine": "mistral_" + mistral_engine.get_mode()}
        from routes.editor_routes import generate_ai_content
        prompt = (
            "Сен профессионал ўзбек илмий муҳаррирсан. Берилган матнни имло, грамматика, синтаксис, "
            "тиниш белгилари бўйича тузат. Структурани сақла. Фақат тузатилган матнни қайтар.\n\n" + text
        )
        out = await generate_ai_content(prompt, prefer="auto")
        result_text = out or text
        # ──── TM learn: save successful AI edit for future reuse ────
        if result_text and result_text != text:
            try:
                import tm_search
                tm_search.save_to_tm(text, result_text, src_lang=f"{lang}_raw", tgt_lang=f"{lang}_improved", source="ai_improve")
            except Exception:
                pass
        return {"improved": result_text, "engine": "cloud"}
    except Exception as e:
        logger.exception("ai-improve failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/grammar-score")
async def grammar_score(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """BERT pseudo-perplexity grammar scorer. 0..1 = how natural the sentence sounds."""
    text = (payload.get("text") or "").strip()
    if not text:
        return {"score": 0.0}
    try:
        import bert_engine
        score = bert_engine.engine.grammar_score(text)
        return {"score": round(float(score), 4)}
    except Exception as e:
        logger.exception("grammar-score failed")
        return {"score": 0.0, "error": str(e)}


@router.post("/translation-quality")
async def translation_quality(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Score translation quality using BERT sentence similarity + LLM critique."""
    src = (payload.get("source") or "").strip()
    tgt = (payload.get("target") or "").strip()
    if not src or not tgt:
        return {"semantic_score": 0.0, "llm_note": ""}
    try:
        import mistral_engine
        if mistral_engine.is_available():
            return await mistral_engine.score_translation(src, tgt)
    except Exception:
        pass
    try:
        import bert_engine
        score = bert_engine.engine.sentence_similarity(src, tgt)
        return {"semantic_score": round(float(score), 4), "llm_note": ""}
    except Exception:
        return {"semantic_score": 0.0, "llm_note": ""}


@router.get("/ai-status")
async def ai_status(current_user: Dict = Depends(get_current_user)):
    """Report which AI engines are available."""
    out = {
        "bert": False,
        "mistral": {"available": False, "mode": "unavailable"},
        "llama": {"available": False, "mode": "unavailable"},
        "russian": {"available": False, "mode": "unavailable"},
        "nllb": {"available": False, "mode": "unavailable"},
        "gemini": False,
        "anthropic": False,
    }
    try:
        import bert_engine
        out["bert"] = bool(bert_engine.engine.initialized)
    except Exception:
        pass
    try:
        import mistral_engine
        out["mistral"] = {"available": mistral_engine.is_available(), "mode": mistral_engine.get_mode(), "model": mistral_engine.MODEL_ID}
    except Exception:
        pass
    try:
        import llama_engine
        out["llama"] = {"available": llama_engine.is_available(), "mode": llama_engine.get_mode(), "model": llama_engine.MODEL_ID}
    except Exception:
        pass
    try:
        import russian_engine
        out["russian"] = {"available": russian_engine.is_available(), "mode": russian_engine.get_mode(), "model": russian_engine.MODEL_ID}
    except Exception:
        pass
    try:
        import translator_engine
        out["nllb"] = {"available": translator_engine.is_available(), "mode": translator_engine.get_mode(), "model": translator_engine.MODEL_ID}
    except Exception:
        pass
    out["gemini"] = bool(os.environ.get("GOOGLE_API_KEY"))
    out["anthropic"] = bool(os.environ.get("ANTHROPIC_API_KEY"))
    return out


@router.post("/synonyms-fast")
async def synonyms_fast(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """
    Hybrid synonyms: fasttext (fast) + BERT cluster (semantic) + DB.
    Falls back gracefully if fasttext unavailable.
    """
    word = (payload.get("word") or "").strip()
    if not word:
        return {"synonyms": []}
    out: List[Dict[str, Any]] = []
    # 1. fasttext
    try:
        import fasttext_engine
        for w, score in fasttext_engine.find_synonyms(word, top_k=15):
            out.append({"word": w, "score": round(score, 3), "source": "fasttext"})
    except Exception:
        pass
    # 2. BERT cluster (if word in classified words)
    try:
        import bert_engine
        if bert_engine.engine.initialized and not out:
            # Use mask to find candidates
            ctx = f"{word} ва [MASK] бир хил маънода."
            preds = bert_engine.engine.predict_mask(ctx, top_k=10, hunspell_filter=True)
            for w in preds:
                if w.lower() != word.lower():
                    out.append({"word": w, "score": 0.7, "source": "bert"})
    except Exception:
        pass
    # 3. Synonyms DB fallback
    try:
        conn = db.connect_db()
        cur = conn.cursor()
        cur.execute("SELECT synonym FROM synonyms WHERE word = ? LIMIT 20", (word,))
        for r in cur.fetchall():
            out.append({"word": r[0], "score": 1.0, "source": "db"})
        conn.close()
    except Exception:
        pass
    # Dedupe
    seen = set()
    final = []
    for s in out:
        if s["word"].lower() not in seen:
            seen.add(s["word"].lower())
            final.append(s)
    return {"synonyms": final[:20]}


@router.post("/normalize-drug")
async def normalize_drug_endpoint(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Find canonical INN/brand for a drug name (cross-script + fuzzy)."""
    name = (payload.get("name") or "").strip()
    if not name:
        return {"input": name, "matches": [], "best": None}
    try:
        import drug_normalizer
        return drug_normalizer.normalize(name, max_results=int(payload.get("max_results", 5)))
    except Exception as e:
        logger.exception("normalize-drug failed")
        return {"input": name, "matches": [], "best": None, "error": str(e)}


@router.get("/training-log/export")
async def export_training_log(kind: str = None, limit: int = 10000, current_user: Dict = Depends(get_current_user)):
    """
    Export collected AI corrections as JSONL for fine-tuning.
    Filter by kind: 'llama_improve', 'llama_translate', 'mistral_improve', 'russian_correct', etc.
    """
    try:
        conn = db.connect_db()
        conn.row_factory = db.sqlite3.Row
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS llm_training_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                kind TEXT, prompt TEXT, response TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        if kind:
            cur.execute("SELECT * FROM llm_training_log WHERE kind = ? ORDER BY created_at DESC LIMIT ?", (kind, limit))
        else:
            cur.execute("SELECT * FROM llm_training_log ORDER BY created_at DESC LIMIT ?", (limit,))
        rows = [dict(r) for r in cur.fetchall()]
        # Also include sayqallash rules as ground-truth pairs
        cur.execute("SELECT wrong_form, correct_form, error_type, lang, frequency FROM sayqallash_rules ORDER BY frequency DESC LIMIT ?", (limit,))
        rules = [{"kind": "sayqallash", "wrong": r[0], "correct": r[1], "error_type": r[2], "lang": r[3], "freq": r[4]} for r in cur.fetchall()]
        conn.close()
        return {
            "training_log": rows,
            "training_log_count": len(rows),
            "sayqallash_rules": rules,
            "sayqallash_count": len(rules),
            "total": len(rows) + len(rules),
        }
    except Exception as e:
        logger.exception("training-log/export failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/training-log/download")
async def download_training_log(kind: str = None, limit: int = 50000, current_user: Dict = Depends(get_current_user)):
    """Download JSONL file of training dataset (for Colab/Kaggle fine-tuning)."""
    try:
        from fastapi.responses import StreamingResponse
        import json as _json
        import io

        conn = db.connect_db()
        conn.row_factory = db.sqlite3.Row
        cur = conn.cursor()
        buf = io.StringIO()

        # Sayqallash pairs
        cur.execute("SELECT wrong_form, correct_form, error_type, lang FROM sayqallash_rules LIMIT ?", (limit,))
        for r in cur.fetchall():
            sample = {
                "instruction": f"Матнни тузат ({r['lang']}):",
                "input": r["wrong_form"],
                "output": r["correct_form"],
                "error_type": r["error_type"],
                "source": "sayqallash",
            }
            buf.write(_json.dumps(sample, ensure_ascii=False) + "\n")

        # LLM log
        try:
            if kind:
                cur.execute("SELECT kind, prompt, response FROM llm_training_log WHERE kind = ? AND LENGTH(prompt) BETWEEN 10 AND 2000 LIMIT ?", (kind, limit))
            else:
                cur.execute("SELECT kind, prompt, response FROM llm_training_log WHERE LENGTH(prompt) BETWEEN 10 AND 2000 LIMIT ?", (limit,))
            for r in cur.fetchall():
                sample = {
                    "instruction": "Илмий-фарма матнни тузат:",
                    "input": r["prompt"][:1500],
                    "output": r["response"][:1500],
                    "source": r["kind"],
                }
                buf.write(_json.dumps(sample, ensure_ascii=False) + "\n")
        except Exception:
            pass

        conn.close()
        content = buf.getvalue().encode("utf-8")
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/x-jsonlines",
            headers={"Content-Disposition": 'attachment; filename="pharma_expert_finetune.jsonl"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/extract-linguistic")
async def extract_linguistic(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """
    Find annotated/disputed/abbreviation words in text.
    Returns spans + 3-language metadata for highlighting + popups.
    """
    text = (payload.get("text") or "").strip()
    if not text:
        return {"annotated": [], "disputed": [], "abbreviations": []}
    out = {"annotated": [], "disputed": [], "abbreviations": []}
    try:
        conn = db.connect_db()
        conn.row_factory = db.sqlite3.Row
        cur = conn.cursor()
        # Build word list from text (lowercase, dedupe)
        import re as _re
        words = set()
        for m in _re.finditer(r"[\w'\u2019\u02bb\u02bc\u2018-]+", text, _re.UNICODE):
            words.add(m.group(0).lower())

        for table, key in (("annotated_words", "annotated"), ("disputed_words", "disputed"), ("abbreviations", "abbreviations")):
            try:
                cur.execute(f"SELECT * FROM {table}")
                cols = [d[0] for d in cur.description]
                for row in cur.fetchall():
                    row_dict = dict(zip(cols, row))
                    # Try common word fields
                    word_val = None
                    for f in ("uz", "word", "short_form", "abbreviation", "uz_lat", "term"):
                        if row_dict.get(f):
                            word_val = row_dict[f].lower()
                            break
                    if not word_val:
                        continue
                    if word_val in words:
                        # Find positions
                        for m in _re.finditer(r"\b" + _re.escape(word_val) + r"\b", text, _re.IGNORECASE):
                            out[key].append({
                                "from": m.start(),
                                "to": m.end(),
                                "text": m.group(0),
                                "data": row_dict,
                            })
            except Exception as e:
                logger.warning(f"[extract-linguistic] {table}: {e}")
        conn.close()
    except Exception as e:
        logger.exception("extract-linguistic failed")
    return out


@router.post("/learn-linguistic")
async def learn_linguistic(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """
    Save user-confirmed annotated/disputed/abbreviation words to DB.
    Payload: {"category": "annotated"|"disputed"|"abbreviations", "items": [{uz, ru, en, definition, ...}]}
    """
    category = payload.get("category", "")
    items = payload.get("items", [])
    table_map = {"annotated": "annotated_words", "disputed": "disputed_words", "abbreviations": "abbreviations"}
    if category not in table_map:
        return {"success": False, "error": "Invalid category"}
    table = table_map[category]
    try:
        conn = db.connect_db()
        cur = conn.cursor()
        # Ensure table exists with basic columns
        cur.execute(f"CREATE TABLE IF NOT EXISTS {table} (id INTEGER PRIMARY KEY AUTOINCREMENT, uz TEXT, ru TEXT, en TEXT, definition TEXT, source TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
        # Try to add missing columns if table existed with different schema
        for col in ("uz", "ru", "en", "definition", "source"):
            try:
                cur.execute(f"ALTER TABLE {table} ADD COLUMN {col} TEXT")
            except Exception:
                pass
        added = 0
        for item in items:
            try:
                cur.execute(
                    f"INSERT INTO {table} (uz, ru, en, definition, source) VALUES (?, ?, ?, ?, ?)",
                    (item.get("uz", ""), item.get("ru", ""), item.get("en", ""), item.get("definition", ""), item.get("source", "user"))
                )
                added += 1
            except Exception as e:
                logger.warning(f"[learn-linguistic] insert: {e}")
        conn.commit()
        conn.close()
        return {"success": True, "added": added}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/tahrirchi/translate")
async def tahrirchi_translate_endpoint(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Translate via tahrirchi/dilmash (native Uzbek model, en/ru/uz/kaa)."""
    text = (payload.get("text") or "").strip()
    src = payload.get("source_lang", "en")
    tgt = payload.get("target_lang", "uz")
    if not text:
        return {"translated": ""}
    try:
        import tahrirchi_engine
        if tahrirchi_engine.is_available():
            result = await tahrirchi_engine.translate_async(text, src, tgt)
            return {"translated": result, "engine": "tahrirchi_" + tahrirchi_engine.get_mode()}
    except Exception as e:
        logger.warning(f"[tahrirchi translate] {e}")
        return {"translated": "", "error": str(e)}
    return {"translated": "", "error": "unavailable"}


@router.post("/tahrirchi/transliterate")
async def tahrirchi_transliterate_endpoint(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Transliterate Uzbek (Latin ↔ Cyrillic) via tahrirchi/dilmash-til."""
    text = (payload.get("text") or "").strip()
    target = payload.get("target", "latin")
    if not text:
        return {"text": ""}
    try:
        import tahrirchi_engine
        result = await tahrirchi_engine.transliterate_async(text, target)
        return {"text": result, "engine": "tahrirchi_til"}
    except Exception as e:
        logger.warning(f"[tahrirchi transliterate] {e}")
        return {"text": text, "error": str(e)}


@router.post("/uzbek/pos-tag")
async def uzbek_pos_tag(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """POS tagging via MaksudSharipov/Uzbek-POS-Tagger-TahrirchiBERT."""
    text = (payload.get("text") or "").strip()
    if not text:
        return {"tags": []}
    try:
        import uzbek_engines
        tags = await uzbek_engines.pos_tag_async(text)
        return {"tags": tags, "engine": "tahrirchi_pos"}
    except Exception as e:
        return {"tags": [], "error": str(e)}


@router.post("/uzbek/correct")
async def uzbek_correct(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Uzbek OCR/transcript corrector via islomov/rubai-corrector."""
    text = (payload.get("text") or "").strip()
    if not text:
        return {"corrected": ""}
    try:
        import uzbek_engines
        corrected = await uzbek_engines.correct_text_async(text)
        return {"corrected": corrected, "engine": "rubai_corrector"}
    except Exception as e:
        return {"corrected": text, "error": str(e)}


@router.get("/uzbek/status")
async def uzbek_status(current_user: Dict = Depends(get_current_user)):
    """Status of advanced Uzbek engines."""
    try:
        import uzbek_engines
        return uzbek_engines.get_status()
    except Exception as e:
        return {"error": str(e)}


@router.get("/style-rules")
async def list_style_rules(category: str = None, current_user: Dict = Depends(get_current_user)):
    """List pharma style guide rules."""
    try:
        conn = db.connect_db()
        conn.row_factory = db.sqlite3.Row
        cur = conn.cursor()
        if category:
            cur.execute("SELECT * FROM style_rules WHERE category = ? ORDER BY rule_id", (category,))
        else:
            cur.execute("SELECT * FROM style_rules ORDER BY rule_id")
        rules = [dict(r) for r in cur.fetchall()]
        conn.close()
        return {"rules": rules, "total": len(rules)}
    except Exception as e:
        return {"rules": [], "error": str(e)}


@router.post("/check-style")
async def check_style(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Check text against pharma style guide rules."""
    text = (payload.get("text") or "").strip()
    if not text:
        return {"violations": []}
    violations = []
    try:
        import re as _re
        conn = db.connect_db()
        conn.row_factory = db.sqlite3.Row
        cur = conn.cursor()
        cur.execute("SELECT * FROM style_rules WHERE pattern != '' AND pattern IS NOT NULL")
        rules = [dict(r) for r in cur.fetchall()]
        for rule in rules:
            try:
                for m in _re.finditer(rule["pattern"], text):
                    violations.append({
                        "rule_id": rule["rule_id"],
                        "category": rule["category"],
                        "severity": rule["severity"],
                        "description": rule["description"],
                        "suggestion": rule["suggestion"],
                        "from": m.start(),
                        "to": m.end(),
                        "matched": m.group(0),
                    })
            except _re.error:
                pass
        conn.close()
    except Exception as e:
        logger.warning(f"[check-style] {e}")
    return {"violations": violations, "total": len(violations)}


@router.post("/decompose-term")
async def decompose_medical_term(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Decompose a medical term into Greek/Latin morphemes."""
    word = (payload.get("word") or "").strip()
    if not word:
        return {"found": False}
    try:
        import compound_morphology
        return compound_morphology.decompose(word)
    except Exception as e:
        return {"found": False, "error": str(e)}


@router.post("/extract-entities")
async def extract_entities_endpoint(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Extract named entities (DRUG, DOSE, PERSON, LOC, ORG) from text."""
    text = (payload.get("text") or "").strip()
    if not text:
        return {"entities": [], "stats": {}}
    try:
        import ner_engine
        ents = await ner_engine.extract_async(text)
        return {"entities": ents, "stats": ner_engine.stats(text)}
    except Exception as e:
        logger.exception("extract-entities failed")
        return {"entities": [], "error": str(e)}


@router.post("/cluster-synonyms")
async def cluster_synonyms(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """Group related words via BERT embedding clustering."""
    words = payload.get("words", [])
    threshold = float(payload.get("threshold", 0.78))
    if not words:
        return {"clusters": []}
    try:
        import bert_engine
        clusters = bert_engine.engine.cluster_words(words, threshold=threshold)
        return {"clusters": clusters}
    except Exception as e:
        return {"clusters": [[w] for w in words], "error": str(e)}


@router.post("/learn-diff")
async def learn_from_diff(payload: Dict[str, Any], current_user: Dict = Depends(get_current_user)):
    """
    Diff original vs corrected text by sentence + word, then add (wrong→correct)
    pairs to sayqallash_rules so the system learns from manual edits.

    Request: { "original": "...", "corrected": "...", "lang": "uz", "source": "tilshunos_edit" }
    """
    original = (payload.get("original") or "").strip()
    corrected = (payload.get("corrected") or "").strip()
    if not original or not corrected or original == corrected:
        return {"learned": 0, "pairs": []}

    lang = payload.get("lang", "uz")
    source = payload.get("source", "tilshunos_edit")

    import difflib

    def split_sents(t: str):
        # Split on . ! ? newline, keep non-empty
        parts = re.split(r'(?<=[\.!\?])\s+|\n+', t)
        return [p.strip() for p in parts if p.strip()]

    def split_words(s: str):
        return re.findall(r"[\w'\u2019\u02bb\u02bc\u2018-]+", s)

    orig_sents = split_sents(original)
    corr_sents = split_sents(corrected)

    sm_sent = difflib.SequenceMatcher(None, orig_sents, corr_sents)
    pairs: List[Dict[str, str]] = []
    for tag, i1, i2, j1, j2 in sm_sent.get_opcodes():
        if tag == 'equal':
            continue
        # Align sentences pairwise within block
        ow = []
        for s in orig_sents[i1:i2]:
            ow += split_words(s)
        cw = []
        for s in corr_sents[j1:j2]:
            cw += split_words(s)
        sm_word = difflib.SequenceMatcher(None, ow, cw)
        for wt, a1, a2, b1, b2 in sm_word.get_opcodes():
            if wt == 'equal':
                continue
            if wt == 'replace' and (a2 - a1) == (b2 - b1):
                for k in range(a2 - a1):
                    w_old, w_new = ow[a1 + k], cw[b1 + k]
                    if w_old.lower() == w_new.lower():
                        continue
                    pairs.append({"wrong": w_old, "correct": w_new})

    learned = 0
    for p in pairs:
        try:
            db.add_sayqallash_rule(
                wrong=p["wrong"],
                correct=p["correct"],
                error_type='S/Spelling',
                context=corrected[:200],
                lang=lang,
                source=source,
            )
            learned += 1
        except Exception as e:
            logger.warning(f"learn-diff add failed for {p}: {e}")

    return {"learned": learned, "pairs": pairs}
