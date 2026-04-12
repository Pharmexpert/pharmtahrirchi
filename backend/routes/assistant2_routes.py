"""
Assistant2 — AI-powered translation, scientific editing, and quality checking.

Uses Claude Sonnet for high-quality pharmaceutical text processing.
System prompts based on O'zbekiston Respublikasi Davlat farmakopeyasi
and European Pharmacopoeia (Ph.Eur.) style and terminology.

Endpoints:
  POST /api/assistant2/translate          — translate text between UZ/RU/EN
  POST /api/assistant2/edit               — scientific editing of pharma text
  POST /api/assistant2/check-translation  — evaluate translation quality (0-100)
  POST /api/assistant2/check-edit         — evaluate edit quality (0-100)
  POST /api/assistant2/upload             — extract text from .txt/.docx files
  POST /api/assistant2/translate-docx     — format-preserving DOCX translation
  POST /api/assistant2/edit-docx          — format-preserving DOCX editing
  POST /api/assistant2/export-pdf         — generate PDF from text
  POST /api/assistant2/text-to-docx       — generate DOCX from text
  POST /api/assistant2/learn              — save user correction for AI learning
"""
import io
import os
import json
import re
import logging
import asyncio
from typing import Dict, Any, Optional, List
from functools import lru_cache

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from auth import get_current_user
from routes.ai_helpers import generate_ai_content, get_anthropic
import db

logger = logging.getLogger("assistant2")
router = APIRouter(prefix="/api/assistant2", tags=["assistant2"])

LANG_NAMES = {"uz": "o'zbek", "ru": "rus", "en": "ingliz"}

# ═══════════════════════════════════════════════════
# Claude Sonnet direct caller
# ═══════════════════════════════════════════════════

async def generate_with_claude(system_prompt: str, user_prompt: str, max_tokens: int = 8192) -> str:
    """Call Claude Sonnet directly with system + user separation.
    Falls back to generate_ai_content() if Anthropic unavailable."""
    client = get_anthropic()
    if client:
        try:
            loop = asyncio.get_event_loop()
            msg = await loop.run_in_executor(
                None,
                lambda: client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=max_tokens,
                    system=system_prompt,
                    messages=[{"role": "user", "content": user_prompt}]
                )
            )
            return msg.content[0].text
        except Exception as e:
            logger.warning(f"[claude-sonnet] failed: {e}, falling back to cascade")
    # Fallback
    return await generate_ai_content(f"{system_prompt}\n\n{user_prompt}", prefer="cloud")


# ═══════════════════════════════════════════════════
# Learned rules loader
# ═══════════════════════════════════════════════════

_learned_rules_cache: Dict[str, Any] = {"rules": "", "ts": 0}

def _load_learned_rules(lang: str = "uz") -> str:
    """Load corrections from sayqallash_rules + hunspell learning to augment prompts.
    Caches for 5 minutes to avoid DB hammering."""
    import time
    now = time.time()
    cache_key = f"{lang}_{now // 300}"  # 5 min buckets
    if _learned_rules_cache.get("key") == cache_key:
        return _learned_rules_cache.get("rules", "")

    rules_text = ""
    try:
        conn = db.get_db()
        cur = conn.cursor()
        # Get top corrections by frequency (sayqallash rules)
        cur.execute("""
            SELECT wrong_form, correct_form, error_type, context
            FROM sayqallash_rules
            WHERE lang = ? AND frequency > 0
            ORDER BY frequency DESC, updated_at DESC
            LIMIT 80
        """, (lang,))
        rows = cur.fetchall()
        if rows:
            rules_parts = []
            for r in rows:
                wrong, correct, etype, ctx = r[0], r[1], r[2] or '', r[3] or ''
                rules_parts.append(f"- \"{wrong}\" → \"{correct}\" ({etype})")
            rules_text = "\n\nO'RGANILGAN QOIDALAR (avvalgi tuzatishlar asosida):\n" + "\n".join(rules_parts)

        _learned_rules_cache["key"] = cache_key
        _learned_rules_cache["rules"] = rules_text
    except Exception as e:
        logger.warning(f"[learned_rules] {e}")

    return rules_text


# ═══════════════════════════════════════════════════
# System prompts — based on DF & EP style
# ═══════════════════════════════════════════════════

TRANSLATE_SYSTEM = """Siz farmatsevtika sohasidagi yuqori malakali mutaxassis tarjimonisiz.

MALAKA:
- O'zbekiston Respublikasi Davlat farmakopeyasi (DF) va Yevropa farmakopeyasi (Ph.Eur.) uslubida ishlaysiz
- Farmatsevtik terminologiyani uch tilda (o'zbek, rus, ingliz) mukammal bilasiz
- INN (Xalqaro patentlanmagan nomlar), ATC kodlar, dozaj shakllari terminlarini bilasiz

TARJIMA QOIDALARI:
1. Farmakopiya terminlarini qat'iy standartlarga mos tarjima qiling:
   - Assay → Миқдорий аниқлаш / Количественное определение
   - Identification → Аниқлаш / Идентификация
   - Related substances → Ёндош аралашмалар / Сопутствующие примеси
   - Dissolution → Эриш / Растворение
   - Loss on drying → Қуритишда йўқотиш / Потеря в массе при высушивании
   - Residue on ignition → Куйдиришдаги қолдиқ / Остаток после прокаливания
   - Uniformity of dosage units → Дозалаш бирликлари бир хиллиги
   - Disintegration → Парчаланиш / Распадаемость
   - Friability → Ишқаланувчанлик / Истираемость
   - Wetting → Намланиш / Смачивание
   - Hardness → Қаттиқлик / Твёрдость
   - Content uniformity → Таркиб бир хиллиги / Однородность содержимого

2. Lotin nomlari (INN, kimyoviy formulalar, ATC kodlar) O'ZGARTIRILMASIN
3. Raqamlar, o'lchov birliklari, formulalar AYNAN saqlansin
4. Jadval va ro'yxat strukturasi saqlansin
5. Farmakopiya maqolasi raqamlari (masalan: 03/2021:20201) o'zgartirilmasin
6. Davlat farmakopeyasi uslubiga mos: rasmiy, ilmiy, aniq

MUHIM — FORMATNI SAQLASH:
- Markdown formatini AYNAN saqlang: sarlavhalar (#, ##, ###), jadvallar (| ... |), qalin (**text**), kursiv (*text*), ro'yxatlar (- yoki 1.)
- Jadval strukturasini buzMANG — | va --- satrlari to'g'ri qaytarilsin
- Formulalar, raqamlar, o'lchov birliklari o'zgartirilmasin
- Bo'sh satrlar va paragraflar saqlansin

FAQAT tarjima matnini qaytaring. Izoh, sharh, qo'shimcha ma'lumot YOZMANG."""

EDIT_SYSTEM = """Siz farmatsevtika sohasidagi ilmiy muharrirsiz.

MALAKA:
- O'zbekiston Respublikasi Davlat farmakopeyasi va Yevropa farmakopeyasi uslubida ishlaysiz
- GxP (GMP, GLP, GCP, GDP) hujjatlari standartlarini bilasiz
- Farmatsevtik terminologiyani uch tilda mukammal bilasiz

TAHRIR MEZONLARI:
1. TERMINOLOGIK ANIQLIK:
   - Farmakopiya terminlari standartga mos bo'lishi (assay, identification, dissolution va h.k.)
   - INN nomlar to'g'ri yozilishi
   - O'lchov birliklari (mg, ml, %, IU) standart formatda

2. USLUBIY IZCHILLIK:
   - Farmakopiya maqolalari uslubi: formal, aniq, qisqa
   - Davlat farmakopeyasi formatiga mos: bob, band, punkt tuzilishi
   - Jadval, formula, rasm havola formati saqlanishi

3. GRAMMATIK TO'G'RILIK:
   - Ilmiy matn grammatikasi
   - O'zbek/Rus/Ingliz tilida terminlar to'g'ri kelishikda

4. FARMATSEVTIK STANDARTLARGA MOSLIK:
   - Sinov usullari tavsiflari to'liq va aniq
   - Mezonlar va normalar raqamli qiymatlari to'g'ri
   - Reagentlar va standart moddalar to'g'ri nomlangan (R - reagent, CRS - standart)

MUHIM — FORMATNI SAQLASH:
- Markdown formatini AYNAN saqlang: sarlavhalar (#, ##, ###), jadvallar (| ... |), qalin (**text**), kursiv (*text*), ro'yxatlar (- yoki 1.)
- Jadval strukturasini buzMANG — | va --- satrlari to'g'ri qaytarilsin
- Formulalar, raqamlar, o'lchov birliklari o'zgartirilmasin
- Bo'sh satrlar va paragraflar saqlansin
- Faqat ilmiy-grammatik tuzatishlar kiriting, strukturani o'zgartirmang

FAQAT tahrirlangan matnni qaytaring. Izoh, sharh, qayd YOZMANG."""

CHECK_TRANSLATION_SYSTEM = """Siz farmatsevtika tarjimasi sifatini baholovchi ekspertsiz.
O'zbekiston Respublikasi Davlat farmakopeyasi va Yevropa farmakopeyasi (Ph.Eur.) standartlarini bilasiz.

Tarjima sifatini quyidagi mezonlar bo'yicha baholang va FAQAT JSON qaytaring:
{
  "umumiy_ball": 0-100,
  "terminologiya": {
    "ball": 0-100,
    "muammolar": ["har bir noto'g'ri tarjima qilingan termin"],
    "tavsiyalar": ["to'g'ri tarjima varianti"]
  },
  "toliqligi": {
    "ball": 0-100,
    "muammolar": ["tushirib qoldirilgan yoki qo'shilgan qismlar"]
  },
  "grammatika": {
    "ball": 0-100,
    "muammolar": ["grammatik xatolar"]
  },
  "uslub": {
    "ball": 0-100,
    "muammolar": ["uslubiy nomuvofiqliklar - farmakopiya uslubiga mos emas"]
  },
  "xulosa": "qisqa xulosa",
  "ijobiy_jihatlar": ["yaxshi tarjima qilingan jihatlar"]
}

BAHOLASH MEZONLARI:
- 90-100: Mukammal farmakopiya tarjimasi, terminlar to'g'ri, uslub mos
- 70-89: Yaxshi, ammo ba'zi terminlar noaniq yoki uslub nomuvofiq
- 50-69: O'rtacha, ko'p terminologik xatolar, mazmun saqlanmagan
- 0-49: Past sifat, jiddiy terminologik va mazmuniy xatolar"""

CHECK_EDIT_SYSTEM = """Siz farmatsevtika ilmiy tahrir sifatini baholovchi ekspertsiz.
O'zbekiston Respublikasi Davlat farmakopeyasi va Yevropa farmakopeyasi standartlarini bilasiz.

Tahrir sifatini baholang va FAQAT JSON qaytaring:
{
  "umumiy_ball": 0-100,
  "ilmiy_aniqlik": {
    "ball": 0-100,
    "izohlar": ["ilmiy jihatdan noto'g'ri yoki noaniq joylar"]
  },
  "ravonlik": {
    "ball": 0-100,
    "izohlar": ["matn ravonligi bilan bog'liq muammolar"]
  },
  "izchillik": {
    "ball": 0-100,
    "izohlar": ["terminologik yoki uslubiy nomuvofiqliklar"]
  },
  "farmatsevtik_standart": {
    "ball": 0-100,
    "izohlar": ["farmakopiya standartlariga mos kelmaydigan joylar"]
  },
  "xulosa": "qisqa xulosa",
  "yaxshilanishlar": ["taklif etiladigan o'zgarishlar"]
}"""


# ═══════════════════════════════════════════════════
# Request models
# ═══════════════════════════════════════════════════

class TranslateRequest(BaseModel):
    text: str
    source_lang: str
    target_langs: List[str]

class EditRequest(BaseModel):
    text: str
    lang: str

class CheckTranslationRequest(BaseModel):
    original: str
    translation: str
    source_lang: str
    target_lang: str

class CheckEditRequest(BaseModel):
    original: str
    edited: str
    lang: str

class LearnRequest(BaseModel):
    wrong: str
    correct: str
    error_type: str = "terminology"
    context: str = ""
    lang: str = "uz"

class ExportPdfRequest(BaseModel):
    text: str
    lang: str = "uz"
    title: str = "Farmatsevtik matn"

class TextToDocxRequest(BaseModel):
    text: str
    title: str = "Farmatsevtik matn"


# ═══════════════════════════════════════════════════
# Text endpoints
# ═══════════════════════════════════════════════════

@router.post("/translate")
async def translate_text(req: TranslateRequest, current_user: Dict = Depends(get_current_user)):
    """Translate pharma text to one or two target languages."""
    if not req.text.strip():
        raise HTTPException(400, "Matn bo'sh")
    if len(req.text) > 50000:
        raise HTTPException(400, "Matn juda uzun (max 50000 belgi)")

    learned = _load_learned_rules(req.source_lang)
    system = TRANSLATE_SYSTEM + learned
    results = {}
    source_name = LANG_NAMES.get(req.source_lang, req.source_lang)

    for target in req.target_langs:
        if target == req.source_lang:
            continue
        target_name = LANG_NAMES.get(target, target)
        user_prompt = f"Manba til: {source_name}\nMaqsad til: {target_name}\n\nMatn:\n{req.text}"
        try:
            result = await generate_with_claude(system, user_prompt)
            results[target] = result.strip() if result else ""
        except Exception as e:
            logger.error(f"[translate] {req.source_lang}->{target} failed: {e}")
            results[target] = f"Xatolik: {str(e)[:200]}"

    return {"translations": results, "source_lang": req.source_lang}


@router.post("/edit")
async def edit_text(req: EditRequest, current_user: Dict = Depends(get_current_user)):
    """Scientific editing of pharma text."""
    if not req.text.strip():
        raise HTTPException(400, "Matn bo'sh")
    if len(req.text) > 50000:
        raise HTTPException(400, "Matn juda uzun (max 50000 belgi)")

    learned = _load_learned_rules(req.lang)
    system = EDIT_SYSTEM + learned
    lang_name = LANG_NAMES.get(req.lang, req.lang)
    user_prompt = f"Til: {lang_name}\n\nMatn:\n{req.text}"

    try:
        result = await generate_with_claude(system, user_prompt)
        return {"edited": result.strip() if result else "", "lang": req.lang}
    except Exception as e:
        logger.error(f"[edit] failed: {e}")
        raise HTTPException(500, f"Tahrir xatoligi: {str(e)[:200]}")


@router.post("/check-translation")
async def check_translation(req: CheckTranslationRequest, current_user: Dict = Depends(get_current_user)):
    """Evaluate translation quality with detailed scoring."""
    if not req.original.strip() or not req.translation.strip():
        raise HTTPException(400, "Original va tarjima matni kerak")

    source_name = LANG_NAMES.get(req.source_lang, req.source_lang)
    target_name = LANG_NAMES.get(req.target_lang, req.target_lang)
    user_prompt = f"Manba til: {source_name}\nMaqsad til: {target_name}\n\nOriginal matn:\n{req.original}\n\nTarjima:\n{req.translation}"

    try:
        result = await generate_with_claude(CHECK_TRANSLATION_SYSTEM, user_prompt)
        match = re.search(r'\{.*\}', result, re.DOTALL)
        if match:
            data = json.loads(match.group())
            return {"result": data}
        return {"result": {"umumiy_ball": 0, "xulosa": result, "raw": True}}
    except Exception as e:
        logger.error(f"[check-translation] failed: {e}")
        raise HTTPException(500, f"Tekshirish xatoligi: {str(e)[:200]}")


@router.post("/check-edit")
async def check_edit(req: CheckEditRequest, current_user: Dict = Depends(get_current_user)):
    """Evaluate edit quality with detailed scoring."""
    if not req.original.strip() or not req.edited.strip():
        raise HTTPException(400, "Dastlabki va tahrirlangan matn kerak")

    lang_name = LANG_NAMES.get(req.lang, req.lang)
    user_prompt = f"Til: {lang_name}\n\nDastlabki matn:\n{req.original}\n\nTahrirlangan matn:\n{req.edited}"

    try:
        result = await generate_with_claude(CHECK_EDIT_SYSTEM, user_prompt)
        match = re.search(r'\{.*\}', result, re.DOTALL)
        if match:
            data = json.loads(match.group())
            return {"result": data}
        return {"result": {"umumiy_ball": 0, "xulosa": result, "raw": True}}
    except Exception as e:
        logger.error(f"[check-edit] failed: {e}")
        raise HTTPException(500, f"Tekshirish xatoligi: {str(e)[:200]}")


# ═══════════════════════════════════════════════════
# File upload
# ═══════════════════════════════════════════════════

@router.post("/upload")
async def upload_file(file: UploadFile = File(...), current_user: Dict = Depends(get_current_user)):
    """Extract text from .txt or .docx file."""
    if not file.filename:
        raise HTTPException(400, "Fayl nomi yo'q")

    name = file.filename.lower()
    raw = await file.read()

    if len(raw) > 10 * 1024 * 1024:
        raise HTTPException(400, "Fayl hajmi 10MB dan oshmasligi kerak")

    if name.endswith('.txt'):
        try:
            text = raw.decode('utf-8')
        except UnicodeDecodeError:
            text = raw.decode('cp1251', errors='replace')
        return {"filename": file.filename, "text": text}

    elif name.endswith('.docx'):
        try:
            md_text = _docx_to_markdown(raw)
            return {"filename": file.filename, "text": md_text}
        except Exception as e:
            raise HTTPException(400, f"DOCX o'qishda xatolik: {str(e)[:200]}")

    else:
        raise HTTPException(400, "Faqat .txt va .docx fayllar qo'llab-quvvatlanadi")


# ═══════════════════════════════════════════════════
# DOCX → Markdown extraction (structure-aware)
# ═══════════════════════════════════════════════════

def _docx_to_markdown(raw_bytes: bytes) -> str:
    """Convert DOCX to Markdown preserving tables, headings, lists, bold/italic."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document(io.BytesIO(raw_bytes))
    md_parts = []

    def runs_to_md(runs) -> str:
        parts = []
        for run in runs:
            text = run.text or ''
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            parts.append(text)
        return ''.join(parts)

    def para_to_md(para) -> str:
        text = runs_to_md(para.runs) if para.runs else para.text
        if not text.strip():
            return ''
        style = (para.style.name or '').lower() if para.style else ''
        if 'heading 1' in style:
            return f"# {text}"
        elif 'heading 2' in style:
            return f"## {text}"
        elif 'heading 3' in style:
            return f"### {text}"
        elif 'heading 4' in style:
            return f"#### {text}"
        elif 'list' in style:
            return f"- {text}"
        return text

    # Body paragraphs and tables
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            # Find matching paragraph object
            for para in doc.paragraphs:
                if para._element is element:
                    line = para_to_md(para)
                    if line:
                        md_parts.append(line)
                    else:
                        md_parts.append('')
                    break
        elif tag == 'tbl':
            # Find matching table
            for table in doc.tables:
                if table._element is element:
                    for ri, row in enumerate(table.rows):
                        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                        md_parts.append('| ' + ' | '.join(cells) + ' |')
                        if ri == 0:
                            md_parts.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
                    md_parts.append('')
                    break

    return '\n'.join(md_parts)


# ═══════════════════════════════════════════════════
# Export endpoints
# ═══════════════════════════════════════════════════

@router.post("/export-pdf")
async def export_pdf(req: ExportPdfRequest, current_user: Dict = Depends(get_current_user)):
    """Generate PDF from text content."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        width, height = A4

        # Try to register a Unicode font
        font_name = "Helvetica"
        for font_path in [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "C:/Windows/Fonts/arial.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont("UniFont", font_path))
                    font_name = "UniFont"
                    break
                except Exception:
                    pass

        # Title
        c.setFont(font_name, 16)
        c.drawString(2 * cm, height - 2 * cm, req.title)

        # Body text
        c.setFont(font_name, 11)
        y = height - 3.5 * cm
        line_height = 14
        max_width = width - 4 * cm

        for line in req.text.split("\n"):
            if y < 2 * cm:
                c.showPage()
                c.setFont(font_name, 11)
                y = height - 2 * cm

            # Simple word-wrap
            words = line.split()
            current_line = ""
            for word in words:
                test = current_line + (" " if current_line else "") + word
                if c.stringWidth(test, font_name, 11) > max_width:
                    if current_line:
                        c.drawString(2 * cm, y, current_line)
                        y -= line_height
                        if y < 2 * cm:
                            c.showPage()
                            c.setFont(font_name, 11)
                            y = height - 2 * cm
                    current_line = word
                else:
                    current_line = test
            if current_line:
                c.drawString(2 * cm, y, current_line)
                y -= line_height
            elif not words:
                y -= line_height  # empty line

        c.save()
        buf.seek(0)

        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{req.title}.pdf"'}
        )
    except ImportError:
        raise HTTPException(500, "reportlab kutubxonasi topilmadi")


@router.post("/text-to-docx")
async def text_to_docx(req: TextToDocxRequest, current_user: Dict = Depends(get_current_user)):
    """Generate a simple DOCX from text."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(req.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Body
        for paragraph_text in req.text.split("\n"):
            p = doc.add_paragraph(paragraph_text)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{req.title}.docx"'}
        )
    except ImportError:
        raise HTTPException(500, "python-docx kutubxonasi topilmadi")


# ═══════════════════════════════════════════════════
# Learning endpoint
# ═══════════════════════════════════════════════════

@router.post("/learn")
async def learn_correction(req: LearnRequest, current_user: Dict = Depends(get_current_user)):
    """Save user correction to improve AI quality."""
    if not req.wrong.strip() or not req.correct.strip():
        raise HTTPException(400, "Xato va to'g'ri variantlar kerak")

    try:
        result = db.learn_from_correction(
            wrong_form=req.wrong.strip(),
            correct_form=req.correct.strip(),
            error_type=req.error_type,
            context=req.context,
            lang=req.lang,
            user_id=current_user.get("id"),
            user_name=current_user.get("name", ""),
        )
        # Clear cache so next request picks up the new rule
        _learned_rules_cache.clear()

        return {"success": True, "rule_id": result.get("rule_id"), "action": result.get("rule_action", "saved")}
    except Exception as e:
        logger.error(f"[learn] failed: {e}")
        raise HTTPException(500, f"Saqlash xatoligi: {str(e)[:200]}")
