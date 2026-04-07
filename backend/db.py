import sqlite3
import json
from typing import List, Dict, Any, Optional
import os
import re
import difflib
import hashlib
import uuid
try:
    from . import transliterate
except ImportError:
    import transliterate
import pickle
import logging
import time
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("db")

try:
    import faiss
    import numpy as np
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False

# Resolve database paths — env vars for Railway, fallback for local dev
BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
# Priority: 1. Env Var, 2. Railway Vol Path (/app/data/), 3. Local (BACKEND_DIR)
IS_RAILWAY = os.getenv("RAILWAY_ENVIRONMENT") or os.path.exists("/app/data")
DATA_DIR = "/app/data" if IS_RAILWAY else BACKEND_DIR

DB_PATH = os.getenv("DB_PATH", os.path.join(DATA_DIR, "pharma_editor.db"))
TAHRIRCHI_DB_PATH = os.getenv("TAHRIRCHI_DB_PATH", os.path.join(DATA_DIR, "tahrirchi.db"))

# Phase 3: FAISS lexicon index over tahrirchi 8.7M words
TAHRIRCHI_FAISS_INDEX_PATH = os.getenv("TAHRIRCHI_FAISS_INDEX", os.path.join(DATA_DIR, "tahrirchi_lexicon.index"))
TAHRIRCHI_FAISS_IDS_PATH = os.getenv("TAHRIRCHI_FAISS_IDS", os.path.join(DATA_DIR, "tahrirchi_lexicon.ids"))

class FaissIndexManager:
    def __init__(self, dimension=768):
        self.dimension = dimension
        self.index = None
        self.rule_ids = [] # map index position to rule ID in sqlite
        if FAISS_AVAILABLE:
            self.index = faiss.IndexFlatIP(dimension)

    def is_ready(self):
        return FAISS_AVAILABLE and self.index is not None and self.index.ntotal > 0

    def add_rule(self, rule_id, vector_bytes):
        if not FAISS_AVAILABLE or self.index is None or not vector_bytes:
            return
        try:
            vec = pickle.loads(vector_bytes)
            if vec is not None:
                # FAISS expects float32
                v_np = vec.astype('float32').reshape(1, -1)
                faiss.normalize_L2(v_np)
                self.index.add(v_np)
                self.rule_ids.append(rule_id)
        except Exception as e:
            logger.error(f"[!] FAISS Add Error: {e}")

    def search(self, query_vec, k=3, threshold=0.92):
        if not self.is_ready() or query_vec is None:
            return []
        try:
            # Query vec as numpy float32
            if not isinstance(query_vec, np.ndarray):
                # Assume it's a torch tensor if not numpy
                query_vec = query_vec.cpu().numpy()
            
            v_np = query_vec.astype('float32').reshape(1, -1)
            faiss.normalize_L2(v_np)
            
            distances, indices = self.index.search(v_np, k)
            
            matches = []
            for dist, idx in zip(distances[0], indices[0]):
                if idx != -1 and idx < len(self.rule_ids) and dist >= threshold:
                    matches.append({
                        "rule_id": self.rule_ids[idx],
                        "score": float(dist)
                    })
            return matches
        except Exception as e:
            logger.error(f"[!] FAISS Search Error: {e}")
            return []

faiss_manager = FaissIndexManager()


# ═══════════════════════════════════════════════════
# Phase 3: Tahrirchi Lexicon (8.7M words) semantic index
# ═══════════════════════════════════════════════════

class TahrirchiLexicon:
    """Semantic search over 8.7M word lexicon using BERT embeddings + FAISS IVF."""

    def __init__(self):
        self.index = None
        self.ids = None
        self.loaded = False
        self.dimension = 768

    def load(self):
        """Load pre-built FAISS index + IDs map from disk."""
        if self.loaded:
            return True
        if not FAISS_AVAILABLE:
            logger.warning("[TahrirchiLex] FAISS not available")
            return False
        if not (os.path.exists(TAHRIRCHI_FAISS_INDEX_PATH) and os.path.exists(TAHRIRCHI_FAISS_IDS_PATH)):
            logger.warning(f"[TahrirchiLex] Index not found at {TAHRIRCHI_FAISS_INDEX_PATH}")
            return False
        try:
            self.index = faiss.read_index(TAHRIRCHI_FAISS_INDEX_PATH)
            with open(TAHRIRCHI_FAISS_IDS_PATH, "rb") as f:
                self.ids = pickle.load(f)
            self.loaded = True
            logger.info(f"[TahrirchiLex] Loaded index with {self.index.ntotal:,} vectors")
            return True
        except Exception as e:
            logger.error(f"[TahrirchiLex] Load error: {e}")
            return False

    def search_similar(self, word: str, k: int = 10, nprobe: int = 16) -> List[Dict[str, Any]]:
        """Find k most semantically similar words in 8.7M lexicon."""
        if not self.load():
            return []
        try:
            import bert_engine
            if not bert_engine.engine.initialized:
                return []

            emb = bert_engine.engine.get_embedding(word.strip().lower(), as_numpy=True)
            if emb is None:
                return []

            # IVF search needs nprobe (trade-off accuracy vs speed)
            if hasattr(self.index, "nprobe"):
                self.index.nprobe = nprobe

            v = emb.astype("float32").reshape(1, -1)
            faiss.normalize_L2(v)
            distances, indices = self.index.search(v, k)

            # Map FAISS indices → word IDs → actual words from tahrirchi.db
            results = []
            if not os.path.exists(TAHRIRCHI_DB_PATH):
                return []
            conn = sqlite3.connect(TAHRIRCHI_DB_PATH)
            cur = conn.cursor()
            for dist, idx in zip(distances[0], indices[0]):
                if idx >= 0 and idx < len(self.ids):
                    word_id = self.ids[idx]
                    cur.execute(
                        "SELECT word, frequency FROM dictionary WHERE id = ?",
                        (word_id,)
                    )
                    row = cur.fetchone()
                    if row:
                        results.append({
                            "word": row[0],
                            "frequency": row[1] if row[1] is not None else 1,
                            "similarity": float(dist),
                        })
            conn.close()
            return results
        except Exception as e:
            logger.error(f"[TahrirchiLex] Search error: {e}")
            return []

    def is_ready(self) -> bool:
        return self.loaded and self.index is not None


tahrirchi_lexicon = TahrirchiLexicon()

def init_faiss_index():
    if not FAISS_AVAILABLE: return
    logger.info("[*] Initializing memory-resident FAISS index...")
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, vector FROM sayqallash_rules WHERE vector IS NOT NULL")
    rows = cursor.fetchall()
    conn.close()
    
    count = 0
    for rid, vec_bin in rows:
        faiss_manager.add_rule(rid, vec_bin)
        count += 1
    logger.info(f"[+] FAISS Index ready with {count} rules.")

def migrate_vectors():
    """Calculate BERT embeddings for any rules missing them (e.g. seed rules)."""
    if not FAISS_AVAILABLE: return
    
    import bert_engine
    if not bert_engine.engine.initialized:
        logger.warning("[!] Cannot migrate: BERT engine not initialized.")
        return

    logger.info("[*] Checking for rules requiring vector migration...")
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, wrong_form, lang FROM sayqallash_rules WHERE vector IS NULL")
    rows = cursor.fetchall()
    
    if not rows:
        logger.info("[+] No rules require migration.")
        conn.close()
        return

    logger.info(f"[*] Migrating {len(rows)} rules (this may take a moment)...")
    count = 0
    for rid, wrong, lang in rows:
        # Re-check BERT initialization in case it fails during loop
        if not bert_engine.engine.initialized:
            time.sleep(1) # Wait a bit and retry
            if not bert_engine.engine.initialized: continue

        emb = bert_engine.engine.get_embedding(wrong.strip().lower(), as_numpy=True)
        if emb is not None:
            vec_bin = pickle.dumps(emb)
            cursor.execute("UPDATE sayqallash_rules SET vector = ? WHERE id = ?", (vec_bin, rid))
            faiss_manager.add_rule(rid, vec_bin)
            count += 1
            
            # Commit every 100 rows to release DB lock for other requests
            if count % 100 == 0:
                conn.commit()
                logger.info(f"[*] Progressive migration: {count}/{len(rows)} rules...")

    conn.commit()
    conn.close()
    logger.info(f"[+] Migration complete. {count} rules updated with vectors.")

class RulesCache:
    def __init__(self):
        self.rules = {} # key: (wrong_form.lower(), lang), value: {correct, type, freq}
        self.last_load = 0
        self.ttl = 300 # cache for 5 minutes

    def load(self):
        try:
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute("SELECT id, wrong_form, correct_form, error_type, lang, frequency FROM sayqallash_rules")
            rows = cursor.fetchall()
            conn.close()

            new_rules = {}
            for r in rows:
                key = (r['wrong_form'].lower().strip(), r['lang'])
                if key not in new_rules or r['frequency'] > new_rules[key]['frequency']:
                    new_rules[key] = {
                        'id': r['id'],
                        'correct': r['correct_form'],
                        'type': r['error_type'],
                        'frequency': r['frequency']
                    }
            self.rules = new_rules
            self.last_load = time.time()
            logger.info(f"[+] RulesCache loaded {len(self.rules)} rules into memory.")
        except Exception as e:
            logger.error(f"[!] RulesCache load error: {e}")

    def get_all(self, lang):
        if time.time() - self.last_load > self.ttl:
            self.load()
        return [
            {'id': v.get('id', 0), 'wrong_form': k[0], 'correct_form': v['correct'], 'error_type': v['type'], 'frequency': v['frequency']}
            for k, v in self.rules.items() if k[1] == lang
        ]

rules_cache = RulesCache()

def connect_db():
    return sqlite3.connect(DB_PATH)

def init_db():
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS projects (
        id TEXT PRIMARY KEY,
        name TEXT,
        specialist_name TEXT,
        user_id TEXT, -- Associated logged-in user
        source_lang TEXT DEFAULT 'English',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        original_filename TEXT,
        file_path TEXT
    )
    ''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS alignments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        sentence_no INTEGER,
        display_no TEXT,
        row_type TEXT DEFAULT 'content',
        en_text TEXT,
        confirmed_ru_text TEXT,
        confirmed_uz_text TEXT,
        text_id TEXT,
        notes TEXT DEFAULT '',
        specialist_name TEXT DEFAULT '',
        user_id TEXT, -- Associated logged-in user
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')

    # Migrate: add display_no/specialist_name/row_type/user_id if not exists
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN display_no TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN specialist_name TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN row_type TEXT DEFAULT 'content'")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN user_id TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN ru_proposed TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN uz_proposed TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN ru_annotations TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN uz_annotations TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN ru_confidence REAL DEFAULT 0")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN uz_confidence REAL DEFAULT 0")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE alignments ADD COLUMN is_pre_polished INTEGER DEFAULT 0")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE projects ADD COLUMN original_filename TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE projects ADD COLUMN file_path TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE projects ADD COLUMN created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE projects ADD COLUMN source_lang TEXT DEFAULT 'English'")
    except Exception: pass

    # ═══════════════════════════════════════════════
    # Sayqallash Rules — self-learning correction DB
    # ═══════════════════════════════════════════════
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS sayqallash_rules (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        wrong_form TEXT NOT NULL,
        correct_form TEXT NOT NULL,
        error_type TEXT DEFAULT 'S/Spelling',
        context TEXT DEFAULT '',
        lang TEXT DEFAULT 'uz',
        frequency INTEGER DEFAULT 1,
        source TEXT DEFAULT 'user_edit',
        vector BLOB, -- Cached BERT embedding
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    # Migrate: add vector column if not exists
    try:
        cursor.execute("ALTER TABLE sayqallash_rules ADD COLUMN vector BLOB")
    except Exception: pass

    # Index for fast lookup
    cursor.execute('''
    CREATE INDEX IF NOT EXISTS idx_sayqallash_wrong ON sayqallash_rules(wrong_form)
    ''')
    cursor.execute('''
    CREATE INDEX IF NOT EXISTS idx_sayqallash_lang ON sayqallash_rules(lang)
    ''')

    # ═══════════════════════════════════════════════
    # User accounts & RBAC
    # ═══════════════════════════════════════════════
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id TEXT PRIMARY KEY,
        email TEXT UNIQUE NOT NULL,
        name TEXT,
        role TEXT DEFAULT 'foydalanuvchi',
        status TEXT DEFAULT 'pending',
        avatar_url TEXT,
        password_hash TEXT,
        salt TEXT,
        department TEXT,
        last_login TIMESTAMP,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    # Migrate: add password columns if not exists
    try:
        cursor.execute("ALTER TABLE users ADD COLUMN password_hash TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE users ADD COLUMN salt TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE users ADD COLUMN department TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE users ADD COLUMN can_edit_db INTEGER DEFAULT 0")
    except Exception: pass
    # Seed admin if empty
    cursor.execute("SELECT COUNT(*) FROM users WHERE email = 'texnopharm@gmail.com'")
    if cursor.fetchone()[0] == 0:
        cursor.execute('''
        INSERT INTO users (id, email, name, role, status, password_hash, salt)
        VALUES ('admin_primary', 'texnopharm@gmail.com', 'Admin Texnopharm', 'admin', 'approved', '8c6976e5b5410415bde908bd4dee15dfb167a9c873fc4bb8a81f6f2ab448a918', 'admin_salt')
        ''')

    # ═══════════════════════════════════════════════
    # Password Resets
    # ═══════════════════════════════════════════════
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS password_resets (
        email TEXT PRIMARY KEY,
        code TEXT NOT NULL,
        expires_at TIMESTAMP NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')

    # ═══════════════════════════════════════════════
    # Phase 11: AI Linguistic Encyclopedia Tables
    # ═══════════════════════════════════════════════
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS annotated_words (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        en TEXT,
        ru TEXT,
        uz TEXT,
        description_en TEXT,
        description_ru TEXT,
        description_uz TEXT,
        source_lang TEXT DEFAULT 'English',
        user_id TEXT,
        text_id TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        modified_by_id TEXT,
        modified_at TIMESTAMP,
        status TEXT DEFAULT 'active'
    )
    ''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS disputed_words (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        en TEXT,
        ru TEXT,
        uz TEXT,
        context_en TEXT,
        context_ru TEXT,
        context_uz TEXT,
        source_lang TEXT DEFAULT 'English',
        user_id TEXT,
        text_id TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        modified_by_id TEXT,
        modified_at TIMESTAMP,
        status TEXT DEFAULT 'active'
    )
    ''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS abbreviations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        short_form TEXT UNIQUE NOT NULL,
        long_en TEXT,
        long_ru TEXT,
        long_uz TEXT,
        source_lang TEXT DEFAULT 'English',
        user_id TEXT,
        text_id TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        modified_by_id TEXT,
        modified_at TIMESTAMP,
        status TEXT DEFAULT 'active'
    )
    ''')

    # Migrate: add source_lang if not exists
    try:
        cursor.execute("ALTER TABLE annotated_words ADD COLUMN source_lang TEXT DEFAULT 'English'")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE disputed_words ADD COLUMN source_lang TEXT DEFAULT 'English'")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE abbreviations ADD COLUMN source_lang TEXT DEFAULT 'English'")
    except Exception: pass

    # ═══════════════════════════════════════════════
    # Phase 12: AI Cache (Optimization)
    # ═══════════════════════════════════════════════
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS ai_cache (
        cache_id TEXT PRIMARY KEY, -- Hash of (lang + prompt_type + text)
        result_json TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_ai_cache_time ON ai_cache(created_at)')

    # ═══════════════════════════════════════════════
    # NEW: Paragraphs Dashboard (History & Audit)
    # ═══════════════════════════════════════════════
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS paragraphs_dashboard (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        en_text TEXT,
        ru_text TEXT,
        uz_text TEXT,
        specialist_name TEXT,
        text_id TEXT,
        action_type TEXT, -- 'AI Polished', 'Manual Edit', 'Verified'
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')

    # ═══════════════════════════════════════════════
    # NEW: Synonyms Database (3-language)
    # ═══════════════════════════════════════════════
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS synonyms (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        word TEXT NOT NULL,
        synonym TEXT NOT NULL,
        lang TEXT NOT NULL DEFAULT 'uz',
        part_of_speech TEXT,
        frequency INTEGER DEFAULT 0,
        probability_scale REAL DEFAULT 0.0,
        source TEXT DEFAULT 'ai',
        author TEXT,
        created_by TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(word, synonym, lang)
    )
    ''')
    try:
        cursor.execute("ALTER TABLE synonyms ADD COLUMN created_by TEXT")
    except Exception: pass
    try:
        cursor.execute("ALTER TABLE projects ADD COLUMN status TEXT DEFAULT 'active'")
    except Exception: pass
    cursor.execute('CREATE INDEX IF NOT EXISTS idx_synonyms_word ON synonyms(word, lang)')

    # Migrate: add new columns to linguistic tables
    for tbl in ['annotated_words', 'disputed_words', 'abbreviations']:
        try:
            cursor.execute(f"ALTER TABLE {tbl} ADD COLUMN text_id TEXT")
        except Exception: pass
        try:
            cursor.execute(f"ALTER TABLE {tbl} ADD COLUMN status TEXT DEFAULT 'active'")
        except Exception: pass

    # Migrate: add modified_by to sayqallash_rules
    try:
        cursor.execute("ALTER TABLE sayqallash_rules ADD COLUMN modified_by TEXT")
    except Exception: pass


    # ═══════════════════════════════════════════════
    # Seed Initial Rules (Pharma Standards)
    # ═══════════════════════════════════════════════
    seed_rules = [
        # Russian Pharma (GF / ГФ)
        ('Растворение', 'Растворимость', 'S/Context', 'ru'),
        ('Растворяемость', 'Растворимость', 'S/Context', 'ru'),
        ('Анализ', 'Количественное определение', 'S/Context', 'ru'), # Assay -> Количественное определение
        ('Идентификация', 'Подлинность', 'S/Context', 'ru'), # Identification -> Подлинность
        ('Сопутствующие вещества', 'Примеси', 'S/Context', 'ru'),
    ]
    for wrong, correct, rtype, l in seed_rules:
        cursor.execute("SELECT 1 FROM sayqallash_rules WHERE wrong_form = ? AND lang = ?", (wrong, l))
        if not cursor.fetchone():
            cursor.execute("INSERT INTO sayqallash_rules (wrong_form, correct_form, error_type, lang, source) VALUES (?, ?, ?, ?, 'seed')",
                           (wrong, correct, rtype, l))

    conn.commit()
    conn.close()

# ═══════════════════════════════════════════════════
# Sayqallash Rules CRUD
# ═══════════════════════════════════════════════════

def add_sayqallash_rule(wrong: str, correct: str, error_type: str = 'S/Spelling',
                         context: str = '', lang: str = 'uz', source: str = 'user_edit'):
    """Add or increment a correction rule with vector caching."""
    if not wrong or not correct or wrong.strip() == correct.strip():
        return
    
    import bert_engine
    import pickle
    vector = None
    if bert_engine.engine.initialized:
        # Calculate vector only once for the rule
        emb = bert_engine.engine.get_embedding(wrong.strip().lower())
        if emb is not None:
            vector = pickle.dumps(emb.cpu().numpy())

    conn = connect_db()
    cursor = conn.cursor()
    # Check if rule exists
    cursor.execute(
        "SELECT id, frequency FROM sayqallash_rules WHERE wrong_form = ? AND correct_form = ? AND lang = ?",
        (wrong.strip(), correct.strip(), lang)
    )
    row = cursor.fetchone()
    if row:
        cursor.execute(
            "UPDATE sayqallash_rules SET frequency = ?, updated_at = CURRENT_TIMESTAMP, context = ?, vector = ? WHERE id = ?",
            (row[1] + 1, context[:200] if context else '', vector, row[0])
        )
        # Update FAISS in-memory if available
        if vector:
            faiss_manager.add_rule(row[0], vector)
    else:
        cursor.execute(
            "INSERT INTO sayqallash_rules (wrong_form, correct_form, error_type, context, lang, source, vector) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (wrong.strip(), correct.strip(), error_type, context[:200] if context else '', lang, source, vector)
        )
        new_id = cursor.lastrowid
        if vector:
            faiss_manager.add_rule(new_id, vector)
    conn.commit()
    conn.close()

def is_word_wrong(word: str, lang: str = 'uz') -> bool:
    """Check if a word is explicitly marked as 'wrong' in any rule using in-memory cache."""
    if not word: return False
    key = (word.lower().strip(), lang)
    if rules_cache.last_load == 0:
        rules_cache.load()
    return key in rules_cache.rules

def get_rules_for_text(text: str, lang: str = 'uz') -> List[Dict]:
    """Find known correction rules that apply to the given text.
    
    IMPORTANT: Does NOT suggest replacing text that is a known correct_form.
    This prevents circular corrections (e.g., marking 'аксарият' as wrong
    when it was previously corrected FROM 'аксарияд' TO 'аксарият').
    """
    if not text:
        return []
    
    # Use in-memory cache instead of database query per call
    rules = rules_cache.get_all(lang)
    # Sort by frequency (most common first) to prioritize certain corrections
    rules.sort(key=lambda x: x['frequency'], reverse=True)

    # Build set of known CORRECT forms (lowercase) — these should NEVER be flagged as wrong
    correct_forms_set = set()
    for rule in rules:
        cf = rule['correct_form']
        if cf:
            correct_forms_set.add(cf.strip().lower())

    found = []
    text_lower = text.lower()
    
    # Pre-calculate embeddings for words in text (potential for semantic search)
    import bert_engine
    word_embeddings = {}
    if bert_engine.engine.initialized:
        for word in set(re.findall(r'\w+', text_lower)):
            if len(word) > 3: # Only embed meaningful words
                word_embeddings[word] = bert_engine.engine.get_embedding(word)

    for rule in rules:
        wrong = rule['wrong_form']
        correct = rule['correct_form']
        if not wrong or not correct:
            continue
        
        wrong_lower = wrong.lower().strip()
        correct_lower = correct.lower().strip()
        
        # CRITICAL: Skip if wrong and correct are same
        if wrong_lower == correct_lower:
            continue
        
        # Exact Match
        start_search = 0
        while True:
            idx = text_lower.find(wrong_lower, start_search)
            if idx == -1:
                break
            
            before_ok = (idx == 0) or not text[idx - 1].isalpha()
            after_ok = (idx + len(wrong) >= len(text)) or not text[idx + len(wrong)].isalpha()
            
            if before_ok and after_ok:
                actual_wrong = text[idx:idx + len(wrong)]
                found.append({
                    'from_index': idx,
                    'to_index': idx + len(wrong),
                    'old_value': actual_wrong,
                    'new_value': correct,
                    'error_type': rule['error_type'],
                    'source': 'rules_db',
                    'frequency': rule['frequency']
                })
            start_search = idx + len(wrong)

    # Semantic (Vector) Match - SCALABLE VERSION (FAISS)
    if bert_engine.engine.initialized:
        import torch
        
        # Create a mapping for quick access to rule data by ID
        rules_by_id = {r['id']: r for r in rules}

        for word, emb in word_embeddings.items():
            if faiss_manager.is_ready():
                # HIGH PERFORMANCE PATH
                matches = faiss_manager.search(emb, k=3, threshold=0.92)
                for m in matches:
                    rule = rules_by_id.get(m['rule_id'])
                    if not rule: continue
                    
                    wrong_lower = rule['wrong_form'].lower().strip()
                    if word == wrong_lower: continue 
                    
                    pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
                    for match_obj in pattern.finditer(text):
                        if not any(f['from_index'] == match_obj.start() for f in found):
                            found.append({
                                'from_index': match_obj.start(),
                                'to_index': match_obj.end(),
                                'old_value': match_obj.group(),
                                'new_value': rule['correct_form'],
                                'error_type': f"FAISS/{rule['error_type']}",
                                'source': 'rules_db_faiss',
                                'frequency': rule['frequency']
                            })
            else:
                # FALLBACK: O(n) loop
                for rule in rules:
                    if not rule['vector']: continue
                    
                    cached_vec = torch.from_numpy(pickle.loads(rule['vector'])).to(emb.device)
                    sim = bert_engine.engine.cosine_similarity(emb, cached_vec)
                    
                    if sim > 0.92:
                        wrong_lower = rule['wrong_form'].lower().strip()
                        if word == wrong_lower: continue 
                        
                        pattern = re.compile(rf'\b{re.escape(word)}\b', re.IGNORECASE)
                        for m in pattern.finditer(text):
                            if not any(f['from_index'] == m.start() for f in found):
                                found.append({
                                    'from_index': m.start(),
                                    'to_index': m.end(),
                                    'old_value': m.group(),
                                    'new_value': rule['correct_form'],
                                    'error_type': f"Semantic/{rule['error_type']}",
                                    'source': 'rules_db_vector',
                                    'frequency': rule['frequency']
                                })
            
    # ═══════════════════════════════════════════════════
    # Dictionary-based spell checking (Tahrirchi 8.7M)
    # Optimized Batch Version
    # ═══════════════════════════════════════════════════
    if lang == 'uz' and os.path.exists(TAHRIRCHI_DB_PATH):
        words = re.findall(r"\w+", text)
        if words:
            # Collect all unique variants to lookup in one batch
            all_variants = set()
            word_to_variants = {}
            for word in words:
                if len(word) < 2: continue
                vars = transliterate.normalize_for_lookup(word)
                all_variants.update(vars)
                word_to_variants[word] = vars
            
            if all_variants:
                dict_conn = sqlite3.connect(TAHRIRCHI_DB_PATH)
                dict_cursor = dict_conn.cursor()
                
                # Use a temporary table or a large IN clause for batch lookup
                # For 8.7M rows, an Index lookup in a batch is much faster than sequential opens
                placeholders = ','.join(['?'] * len(all_variants))
                
                # Check for FTS5 table first, fallback to standard B-tree
                try:
                    # FTS5 for prefix/fuzzy matches if needed, but for exact 'IN' lookup, standard index is fine
                    dict_cursor.execute(f"SELECT word FROM dictionary WHERE word IN ({placeholders})", list(all_variants))
                    found_variants = {row[0] for row in dict_cursor.fetchall()}
                except Exception:
                    found_variants = set()
                
                dict_conn.close()
                
                for word, vars in word_to_variants.items():
                    if not any(v in found_variants for v in vars):
                        # Word not in 8.7M dictionary -> mark as potential typo
                        idx = text.lower().find(word.lower())
                        if idx != -1:
                            found.append({
                                'from_index': idx,
                                'to_index': idx + len(word),
                                'old_value': word,
                                'new_value': '[Луғатда топилмади]',
                                'error_type': 'S/Spelling',
                                'source': 'tahrirchi_dict',
                                'frequency': 0
                            })

def search_dictionary(query: str, limit: int = 10) -> List[str]:
    """Prefix search in the 8.7M word dictionary using FTS5."""
    if not query or len(query) < 2 or not os.path.exists(TAHRIRCHI_DB_PATH):
        return []
    try:
        conn = sqlite3.connect(TAHRIRCHI_DB_PATH)
        cursor = conn.cursor()
        # Use MATCH for FTS5 prefix search
        cursor.execute(
            "SELECT word FROM dictionary_fts WHERE word MATCH ? ORDER BY rank LIMIT ?",
            (f"{query.strip()}*", limit)
        )
        rows = cursor.fetchall()
        conn.close()
        return [r[0] for r in rows]
    except Exception as e:
        print(f"Dictionary search error: {e}")
        return []

    return found

def get_all_rules(lang: str = 'uz', limit: int = 500) -> List[Dict]:
    """Get all rules for a language, excluding binary vector data."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    # Explicitly exclude 'vector' as it contains binary data that breaks JSON encoding
    cursor.execute(
        "SELECT id, wrong_form, correct_form, error_type, context, lang, frequency, created_at, updated_at FROM sayqallash_rules WHERE lang = ? ORDER BY frequency DESC LIMIT ?",
        (lang, limit)
    )
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_rules_count(lang: str = 'uz') -> int:
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM sayqallash_rules WHERE lang = ?", (lang,))
    count = cursor.fetchone()[0]
    conn.close()
    return count

def delete_sayqallash_rule(rule_id: int):
    """Delete a rule by ID."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM sayqallash_rules WHERE id = ?", (rule_id,))
    conn.commit()
    conn.close()

def update_sayqallash_rule(rule_id: int, data: Dict[str, Any]):
    """Update an existing rule."""
    if not data:
        return
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Dynamically build the UPDATE statement
    fields = []
    values = []
    for k, v in data.items():
        if k in ['wrong_form', 'correct_form', 'error_type', 'context', 'lang', 'frequency']:
            fields.append(f"{k} = ?")
            values.append(v)
    
    if fields:
        fields.append("updated_at = CURRENT_TIMESTAMP")
        values.append(rule_id)
        query = f"UPDATE sayqallash_rules SET {', '.join(fields)} WHERE id = ?"
        cursor.execute(query, tuple(values))
        conn.commit()
    conn.close()

# ═══════════════════════════════════════════════════
# Auto-diff: extract correction rules from V1 vs Proposed
# ═══════════════════════════════════════════════════

def extract_diff_rules(v1_text: str, proposed_text: str, lang: str = 'uz', en_context: str = '') -> List[Dict]:
    """
    Compare V1 (original) with Proposed (corrected) and extract word-level differences
    as correction rules for the self-learning database.
    """
    if not v1_text or not proposed_text or v1_text.strip() == proposed_text.strip():
        return []

    v1_words = v1_text.split()
    proposed_words = proposed_text.split()

    sm = difflib.SequenceMatcher(None, v1_words, proposed_words)
    rules = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'replace':
            wrong = ' '.join(v1_words[i1:i2])
            correct = ' '.join(proposed_words[j1:j2])
            if wrong != correct:
                # Determine error type
                error_type = classify_error(wrong, correct)
                rules.append({
                    'wrong_form': wrong,
                    'correct_form': correct,
                    'error_type': error_type,
                    'context': en_context[:200] if en_context else '',
                    'lang': lang
                })
        elif tag == 'delete':
            wrong = ' '.join(v1_words[i1:i2])
            rules.append({
                'wrong_form': wrong,
                'correct_form': '',
                'error_type': 'F/Clarity',
                'context': en_context[:200] if en_context else '',
                'lang': lang
            })
        elif tag == 'insert':
            correct = ' '.join(proposed_words[j1:j2])
            # Context: words around insertion point
            ctx_word = v1_words[i1-1] if i1 > 0 else ''
            rules.append({
                'wrong_form': f'[{ctx_word}]_missing',
                'correct_form': correct,
                'error_type': 'G/Other',
                'context': en_context[:200] if en_context else '',
                'lang': lang
            })
    return rules

def classify_error(wrong: str, correct: str) -> str:
    """Classify the type of error based on the difference."""
    w, c = wrong.lower(), correct.lower()

    # Check for pure case difference
    if w == c:
        return 'S/LowerUpper'

    # Check spelling (similar words, Levenshtein distance <= 2)
    if len(wrong) > 2 and len(correct) > 2:
        common = sum(1 for a, b in zip(w, c) if a == b)
        similarity = common / max(len(w), len(c))
        if similarity > 0.6:
            return 'S/Spelling'

    # Check merge/split
    if ' ' in wrong and ' ' not in correct:
        return 'G/Merge'
    if ' ' not in wrong and ' ' in correct:
        return 'G/Split'

    # Punctuation
    if re.sub(r'[^\w\s]', '', w) == re.sub(r'[^\w\s]', '', c):
        return 'Punctuation'

    return 'S/Context'

def generate_diff_notes(v1_text: str, proposed_text: str, lang: str = 'uz') -> str:
    """Generate human-readable notes about differences between V1 and Proposed."""
    if not v1_text or not proposed_text or v1_text.strip() == proposed_text.strip():
        return ''

    rules = extract_diff_rules(v1_text, proposed_text, lang)
    if not rules:
        return ''

    lang_label = 'UZ' if lang == 'uz' else 'RU'
    notes_parts = [f"[{lang_label} ўзгаришлар]:"]
    for r in rules:
        if r['correct_form']:
            notes_parts.append(f"  «{r['wrong_form']}» → «{r['correct_form']}» [{r['error_type']}]")
        else:
            notes_parts.append(f"  «{r['wrong_form']}» ўчирилди [{r['error_type']}]")
    return '\n'.join(notes_parts)

def save_corrections_as_rules(item: Dict[str, Any]):
    """Extract and save correction rules from a row's V1 vs Proposed differences."""
    # Uzbek corrections
    uz_v1 = item.get('uz_v1', '')
    uz_proposed = item.get('uz_proposed', '')
    en_context = item.get('en', '')

    if uz_v1 and uz_proposed and uz_v1.strip() != uz_proposed.strip():
        rules = extract_diff_rules(uz_v1, uz_proposed, 'uz', en_context)
        for r in rules:
            if r['correct_form']:  # Only save actual corrections
                add_sayqallash_rule(
                    r['wrong_form'], r['correct_form'], r['error_type'],
                    r['context'], r['lang'], 'user_edit'
                )

    # Russian corrections
    ru_v1 = item.get('ru_v1', '')
    ru_proposed = item.get('ru_proposed', '')
    if ru_v1 and ru_proposed and ru_v1.strip() != ru_proposed.strip():
        rules = extract_diff_rules(ru_v1, ru_proposed, 'ru', en_context)
        for r in rules:
            if r['correct_form']:
                add_sayqallash_rule(
                    r['wrong_form'], r['correct_form'], r['error_type'],
                    r['context'], r['lang'], 'user_edit'
                )


def save_alignments(project_id: str, alignments: List[Dict[str, Any]], user_id: str = None):
    conn = connect_db()
    cursor = conn.cursor()
    # Clear old alignments for this project
    cursor.execute("DELETE FROM alignments WHERE text_id = ?", (project_id,))
    
    for row in alignments:
        # Save to alignments table
        cursor.execute('''
            INSERT INTO alignments (sentence_no, display_no, row_type, en_text, confirmed_ru_text, confirmed_uz_text, text_id, notes, specialist_name, user_id)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            row.get('sentence_no'),
            row.get('display_no'),
            row.get('row_type', 'content'),
            row.get('en_text'),
            row.get('confirmed_ru_text'),
            row.get('confirmed_uz_text'),
            project_id,
            row.get('notes', ''),
            row.get('specialist_name', 'Aniqlanmagan'),
            user_id
        ))
        
        # Batch record confirmed content rows in Paragraphs Dashboard
        if row.get('row_type', 'content') == 'content' and (row.get('confirmed_ru_text') or row.get('confirmed_uz_text')):
            try:
                record_dashboard_entry(
                    en=row.get('en_text', ''),
                    ru=row.get('confirmed_ru_text', ''),
                    uz=row.get('confirmed_uz_text', ''),
                    specialist=row.get('specialist_name', 'Aniqlanmagan'),
                    text_id=project_id,
                    action_type='Batch Save'
                )
            except Exception as e:
                print(f"Batch dashboard record error: {e}")

    conn.commit()
    conn.close()

    
    # Update project metadata
    if alignments:
        update_project_metadata(project_id, alignments[0].get("specialist_name", ""))

def save_project_polishing_summary(project_id: str, summary: Dict):
    """Save the aggregate results of a whole-document polishing run."""
    try:
        conn = connect_db()
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE projects SET batch_polishing_summary = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (json.dumps(summary), project_id)
        )
        conn.commit()
        conn.close()
    except Exception as e:
        logger.error(f"Error saving polishing summary: {e}")

def get_project_polishing_summary(project_id: str) -> Optional[Dict]:
    """Retrieve the last recorded polishing summary for a project."""
    try:
        conn = connect_db()
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT batch_polishing_summary FROM projects WHERE id = ?", (project_id,))
        row = cursor.fetchone()
        conn.close()
        if row and row['batch_polishing_summary']:
            return json.loads(row['batch_polishing_summary'])
    except Exception as e:
        logger.error(f"Error getting polishing summary: {e}")
    return None

def update_project_metadata(text_id: str, specialist: str = "", user_id: str = None):
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM projects WHERE id = ?", (text_id,))
    if cursor.fetchone():
        if user_id:
            cursor.execute("UPDATE projects SET specialist_name = ?, user_id = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (specialist, user_id, text_id))
        else:
            cursor.execute("UPDATE projects SET specialist_name = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (specialist, text_id))
    else:
        cursor.execute("INSERT INTO projects (id, name, specialist_name, user_id) VALUES (?, ?, ?, ?)", (text_id, f"Project {text_id}", specialist, user_id))
    conn.commit()
    conn.close()

def add_project(project_id: str, name: str, specialist_name: str = "Aniqlanmagan", user_id: str = None):
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO projects (id, name, specialist_name, user_id) VALUES (?, ?, ?, ?)",
        (project_id, name, specialist_name, user_id)
    )
    conn.commit()
    conn.close()

def list_projects() -> List[Dict[str, Any]]:
    conn = connect_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    # Join with users to get the real name and email
    cursor.execute('''
        SELECT p.*, p.id as text_id, u.name as user_full_name, u.email as user_email
        FROM projects p
        LEFT JOIN users u ON p.user_id = u.id
        ORDER BY p.updated_at DESC
    ''')
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]

def delete_project(text_id: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM projects WHERE id = ?", (text_id,))
    cursor.execute("DELETE FROM alignments WHERE text_id = ?", (text_id,))
    conn.commit()
    conn.close()

def get_alignments_by_text_id(text_id: str) -> List[Dict[str, Any]]:
    """Retrieve all rows for a project."""
    conn = connect_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM alignments WHERE text_id = ? ORDER BY sentence_no ASC", (text_id,))
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]

def update_alignment_ai_result(row_id: int, lang: str, corrected_text: str, annotations: List[Dict], confidence: float):
    """Update a specific row with AI results (used by background task)."""
    conn = connect_db()
    cursor = conn.cursor()
    if lang == 'ru':
        cursor.execute('''
            UPDATE alignments SET ru_proposed = ?, ru_annotations = ?, ru_confidence = ?, is_pre_polished = 1 
            WHERE id = ?
        ''', (corrected_text, json.dumps(annotations), confidence, row_id))
    else:
        cursor.execute('''
            UPDATE alignments SET uz_proposed = ?, uz_annotations = ?, uz_confidence = ?, is_pre_polished = 1
            WHERE id = ?
        ''', (corrected_text, json.dumps(annotations), confidence, row_id))
    conn.commit()
    conn.close()

def save_single_row(item: Dict[str, Any], user_id: str = None) -> int:
    """Save or update a single row. Returns the real DB id."""
    conn = connect_db()
    cursor = conn.cursor()
    text_id = item.get("text_id", "default")
    sentence_no = item.get("sentence_no", 0)

    # Auto-extract correction rules
    save_corrections_as_rules(item)

    if sentence_no and sentence_no > 0:
        # Check by sentence_no + text_id
        cursor.execute(
            "SELECT id FROM alignments WHERE text_id = ? AND sentence_no = ?",
            (text_id, sentence_no)
        )
        row = cursor.fetchone()
        if row:
            cursor.execute('''
            UPDATE alignments SET
                en_text = ?, confirmed_ru_text = ?, confirmed_uz_text = ?,
                notes = ?, display_no = ?, specialist_name = ?, user_id = ?
            WHERE id = ?
            ''', (
                item.get("en", ""),
                item.get("ru_proposed", ""),
                item.get("uz_proposed", ""),
                item.get("notes", ""),
                item.get("display_no", str(sentence_no)),
                item.get("specialist_name", ""),
                user_id,
                row[0]
            ))
            conn.commit()
            conn.close()
            # Record in Paragraphs Dashboard (History)
            try:
                record_dashboard_entry(
                    en=item.get("en", ""),
                    ru=item.get("ru_proposed", ""),
                    uz=item.get("uz_proposed", ""),
                    specialist=item.get("specialist_name", "Aniqlanmagan"),
                    text_id=text_id,
                    action_type=item.get("action_type", "Manual Edit")
                )
            except Exception as e:
                print(f"Single row dashboard record error: {e}")
            return row[0]

    # New row — INSERT
    cursor.execute('''
    INSERT INTO alignments (sentence_no, display_no, en_text, confirmed_ru_text, confirmed_uz_text, text_id, notes, specialist_name, user_id)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        sentence_no if sentence_no and sentence_no > 0 else None,
        item.get("display_no", ""),
        item.get("en", ""),
        item.get("ru_proposed", ""),
        item.get("uz_proposed", ""),
        text_id,
        item.get("notes", ""),
        item.get("specialist_name", ""),
        user_id
    ))
    new_id = cursor.lastrowid
    # Store DB id as sentence_no for future updates
    cursor.execute("UPDATE alignments SET sentence_no = ? WHERE id = ?", (new_id, new_id))
    conn.commit()
    conn.close()
    
    # Record in Paragraphs Dashboard (History)
    try:
        record_dashboard_entry(
            en=item.get("en", ""),
            ru=item.get("ru_proposed", ""),
            uz=item.get("uz_proposed", ""),
            specialist=item.get("specialist_name", "Aniqlanmagan"),
            text_id=text_id,
            action_type=item.get("action_type", "Manual Edit")
        )
    except Exception as e:
        print(f"Single row dashboard record error: {e}")
        
    return new_id

def delete_row(sentence_no: int, text_id: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute(
        "DELETE FROM alignments WHERE text_id = ? AND sentence_no = ?",
        (text_id, sentence_no)
    )
    conn.commit()
    conn.close()

def get_history(text_id: str):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute(
        "SELECT * FROM alignments WHERE text_id = ? ORDER BY sentence_no ASC",
        (text_id,)
    )
    rows = cursor.fetchall()
    conn.close()
    
    # Map DB column names to frontend field names
    result = []
    for r in rows:
        d = dict(r)
        result.append({
            'type': d.get('row_type', 'content'),
            'sentence_no': d.get('sentence_no', 0),
            'display_no': d.get('display_no', str(d.get('sentence_no', ''))),
            'en': d.get('en_text', ''),
            'ru_v1': d.get('confirmed_ru_text', ''),
            'ru_proposed': d.get('confirmed_ru_text', ''),
            'uz_v1': d.get('confirmed_uz_text', ''),
            'uz_proposed': d.get('confirmed_uz_text', ''),
            'text_id': d.get('text_id', text_id),
            'notes': d.get('notes', ''),
            'status': 'aligned',
            'specialist_name': d.get('specialist_name', '')
        })
    return result


# ═══════════════════════════════════════════════════
# User Management
# ═══════════════════════════════════════════════════

def get_user_by_email(email: str) -> Optional[Dict[str, Any]]:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE email = ?", (email.lower().strip(),))
    row = cursor.fetchone()
    conn.close()
    return dict(row) if row else None

def get_user_by_id(user_id: str) -> Optional[Dict[str, Any]]:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE id = ?", (user_id,))
    row = cursor.fetchone()
    conn.close()
    return dict(row) if row else None

def create_reset_code(email: str, code: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    # Expire in 15 minutes
    expires_at = time.time() + (15 * 60)
    cursor.execute('''
    INSERT OR REPLACE INTO password_resets (email, code, expires_at)
    VALUES (?, ?, ?)
    ''', (email.lower().strip(), code, expires_at))
    conn.commit()
    conn.close()

def get_reset_code(email: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT code, expires_at FROM password_resets WHERE email = ?", (email.lower().strip(),))
    row = cursor.fetchone()
    conn.close()
    if row:
        code, expires_at = row
        if time.time() < expires_at:
            return code
    return None

def delete_reset_code(email: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM password_resets WHERE email = ?", (email.lower().strip(),))
    conn.commit()
    conn.close()

def update_user_password(email: str, new_password: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    salt = uuid.uuid4().hex
    password_hash = hashlib.sha256((new_password + salt).encode()).hexdigest()
    cursor.execute('''
    UPDATE users SET password_hash = ?, salt = ? WHERE email = ?
    ''', (password_hash, salt, email.lower().strip()))
    conn.commit()
    conn.close()

def create_user(user_id: str, email: str, name: str, department: Optional[str] = None, avatar_url: Optional[str] = None, password: Optional[str] = None, auto_approve: bool = False):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    password_hash = None
    salt = None
    if password:
        salt = uuid.uuid4().hex
        password_hash = hashlib.sha256((password + salt).encode()).hexdigest()

    # Google OAuth users get auto-approved, manual registrations stay pending
    status = 'approved' if auto_approve else 'pending'
    
    cursor.execute('''
    INSERT INTO users (id, email, name, department, avatar_url, status, password_hash, salt)
    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''', (user_id, email.lower().strip(), name, department, avatar_url, status, password_hash, salt))
    conn.commit()
    conn.close()

def verify_password(email: str, password: str) -> bool:
    user = get_user_by_email(email)
    if not user:
        return False
        
    stored_hash = user.get('password_hash')
    salt = user.get('salt')
    role = user.get('role', 'user')

    # Primary Admin Passwordless Bypass (If no password is set)
    if not stored_hash and not salt and role == 'admin':
        return True

    if not stored_hash or not salt:
        return False

    # If admin with legacy hardcoded hash
    if email == 'texnopharm@gmail.com' and stored_hash == '8c6976e5b5410415bde908bd4dee15dfb167a9c873fc4bb8a81f6f2ab448a918':
        return password == 'admin123' or hashlib.sha256((password + salt).encode()).hexdigest() == stored_hash

    return hashlib.sha256((password + salt).encode()).hexdigest() == stored_hash

def update_user_login(user_id: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE id = ?", (user_id,))
    conn.commit()
    conn.close()

def list_all_users() -> List[Dict[str, Any]]:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users ORDER BY created_at DESC")
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]

def update_user_role(user_id: str, role: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("UPDATE users SET role = ? WHERE id = ?", (role, user_id))
    conn.commit()
    conn.close()

def reorder_alignments(text_id: str):
    """Ensure sentence_no and display_no are sequential for a project."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM alignments WHERE text_id = ? ORDER BY id ASC", (text_id,))
    rows = cursor.fetchall()
    for idx, row in enumerate(rows, 1):
        cursor.execute("UPDATE alignments SET sentence_no = ?, display_no = ? WHERE id = ?", (idx, str(idx), row[0]))
    conn.commit()
    conn.close()

def delete_alignment(alignment_id: int):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM alignments WHERE id = ?", (alignment_id,))
    conn.commit()
    conn.close()

def get_project_file_path(project_id: str) -> Optional[str]:
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT file_path FROM projects WHERE id = ?", (project_id,))
    row = cursor.fetchone()
    conn.close()
    return row[0] if row else None

def delete_project(project_id: str):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM projects WHERE id = ?", (project_id,))
    cursor.execute("DELETE FROM alignments WHERE text_id = ?", (project_id,))
    conn.commit()
    conn.close()


def get_unique_specialists() -> List[str]:
    """Get unique specialist names from both alignments and projects tables."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    # Union names from both tables for robustness
    query = """
        SELECT DISTINCT specialist_name FROM alignments WHERE specialist_name IS NOT NULL AND specialist_name != ''
        UNION
        SELECT DISTINCT specialist_name FROM projects WHERE specialist_name IS NOT NULL AND specialist_name != ''
        ORDER BY specialist_name
    """
    cursor.execute(query)
    rows = cursor.fetchall()
    conn.close()
    return [r[0] for r in rows]


def update_user_status(user_id: str, status: str):
    """Update user approval status (pending/approved/rejected)."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("UPDATE users SET status = ? WHERE id = ?", (status, user_id))
    conn.commit()
    conn.close()

def update_user_profile(user_id: str, data: Dict[str, Any]):
    """Update personal profile information."""
    conn = connect_db()
    cursor = conn.cursor()
    fields = []
    params = []
    
    if 'name' in data:
        fields.append("name = ?")
        params.append(data['name'])
    if 'email' in data:
        fields.append("email = ?")
        params.append(data['email'].lower().strip())
    if 'department' in data:
        fields.append("department = ?")
        params.append(data['department'])
    if 'password' in data and data['password']:
        salt = uuid.uuid4().hex
        password_hash = hashlib.sha256((data['password'] + salt).encode()).hexdigest()
        fields.append("password_hash = ?")
        params.append(password_hash)
        fields.append("salt = ?")
        params.append(salt)
        
    if fields:
        params.append(user_id)
        query = f"UPDATE users SET {', '.join(fields)} WHERE id = ?"
        cursor.execute(query, tuple(params))
        conn.commit()
    conn.close()
# ═══════════════════════════════════════════════
# AI Cache Helpers
# ═══════════════════════════════════════════════

def get_ai_cache(text: str, lang: str, prompt_type: str = "sayqallash") -> Optional[Dict]:
    """Retrieve AI result from cache to avoid redundant API calls."""
    if not text: return None
    # Simple hash of identifying factors
    key_str = f"{lang}:{prompt_type}:{text.strip()}"
    cache_id = hashlib.sha256(key_str.encode()).hexdigest()
    
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("SELECT result_json FROM ai_cache WHERE cache_id = ?", (cache_id,))
    row = cursor.fetchone()
    conn.close()
    
    if row:
        try:
            return json.loads(row[0])
        except Exception:
            return None
    return None

def set_ai_cache(text: str, lang: str, result_dict: Dict, prompt_type: str = "sayqallash"):
    """Store AI result in cache."""
    if not text or not result_dict: return
    key_str = f"{lang}:{prompt_type}:{text.strip()}"
    cache_id = hashlib.sha256(key_str.encode()).hexdigest()
    
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("INSERT OR REPLACE INTO ai_cache (cache_id, result_json) VALUES (?, ?)", 
                   (cache_id, json.dumps(result_dict)))
    conn.commit()
    conn.close()

# ═══════════════════════════════════════════════
# NEW: Paragraphs Dashboard CRUD
# ═══════════════════════════════════════════════

def record_dashboard_entry(en: str, ru: str, uz: str, specialist: str, text_id: str, action_type: str):
    """Record a new entry in the Paragraphs Dashboard history (Audit Trail).
    Skips 'Verified' / 'All rows verified' entries — these are bulk confirmations, not real edits.
    """
    # Filter out verification-only entries
    SKIP_TYPES = {'verified', 'all rows verified', 'barcha qatorlar tasdiqlandi'}
    if action_type and action_type.strip().lower() in SKIP_TYPES:
        return

    conn = connect_db()
    cursor = conn.cursor()

    cursor.execute('''
        INSERT INTO paragraphs_dashboard (en_text, ru_text, uz_text, specialist_name, text_id, action_type)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (en, ru, uz, specialist, text_id, action_type))

    conn.commit()
    conn.close()

def sync_paragraphs_to_alignments(text_id: str, en: str, ru: str, uz: str):
    """Synchronize an edit from Paragraphs Dashboard back to the primary alignments table."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute('''
        UPDATE alignments 
        SET confirmed_ru_text = ?, confirmed_uz_text = ?, updated_at = CURRENT_TIMESTAMP
        WHERE text_id = ? AND en_text = ?
    ''', (ru, uz, text_id, en))
    
    # Update project timestamp
    cursor.execute("UPDATE projects SET updated_at = CURRENT_TIMESTAMP WHERE id = ?", (text_id,))
    
    conn.commit()
    conn.close()

def get_dashboard_entries():
    """Retrieve all history from the Paragraphs Dashboard.
    Excludes 'Verified' entries — bulk confirmations are not real edits.
    """
    conn = connect_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("""
        SELECT * FROM paragraphs_dashboard
        WHERE LOWER(action_type) NOT IN ('verified', 'all rows verified', 'barcha qatorlar tasdiqlandi')
        ORDER BY created_at DESC
    """)
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]


def list_uploaded_files(current_user_id: str = None, is_admin: bool = False, current_user_name: str = "") -> List[Dict[str, Any]]:
    """List uploaded files. If is_admin, return all with owner info.
    Otherwise only return files in the user's own folder uploads/users/<safe_name>/..."""
    import os, re as _re
    _IS_RAILWAY = os.getenv("RAILWAY_ENVIRONMENT") or os.path.exists("/app/data")
    _DATA_DIR = os.getenv("DATA_DIR", "/app/data" if _IS_RAILWAY else os.path.dirname(os.path.abspath(__file__)))
    UPLOADS_DIR = os.path.join(_DATA_DIR, "uploads")
    os.makedirs(UPLOADS_DIR, exist_ok=True)

    # Compute user's safe folder name (must match _safe_user_folder in upload_routes)
    def _safe_folder_name_local(name: str, fallback: str) -> str:
        if not name:
            return f"user_{fallback}" if fallback else "user_unknown"
        cleaned = _re.sub(r"[^A-Za-z0-9_]+", "_", name).strip("_")
        return cleaned or (f"user_{fallback}" if fallback else "user_unknown")

    user_safe = _safe_folder_name_local(current_user_name, current_user_id or "unknown") if current_user_id else None

    files = []
    for root, dirs, fnames in os.walk(UPLOADS_DIR):
        rel_dir = os.path.relpath(root, UPLOADS_DIR)
        if rel_dir == ".":
            rel_dir = ""
        for fname in fnames:
            fpath = os.path.join(root, fname)
            try:
                stat = os.stat(fpath)
            except OSError:
                continue
            rel_path = os.path.join(rel_dir, fname).replace("\\", "/") if rel_dir else fname
            files.append({
                "filename": rel_path,
                "path": fpath,
                "size": stat.st_size,
                "modified_at": stat.st_mtime,
                "extension": os.path.splitext(fname)[1].lower(),
                "folder": rel_dir.replace("\\", "/"),
            })

    # Enrich with project + owner info
    conn = connect_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("""
        SELECT p.id, p.name, p.specialist_name, p.original_filename, p.file_path, p.user_id,
               p.created_at, p.updated_at,
               u.name as user_full_name
        FROM projects p LEFT JOIN users u ON p.user_id = u.id
        WHERE p.file_path IS NOT NULL
    """)
    projects = {dict(r)["file_path"]: dict(r) for r in cursor.fetchall()}
    conn.close()

    def _safe_folder_name(name: str, fallback: str) -> str:
        import re as _re
        if not name:
            return fallback or "legacy"
        cleaned = _re.sub(r"[^A-Za-z0-9_]+", "_", name).strip("_")
        return cleaned or (fallback or "legacy")

    enriched = []
    for f in files:
        proj = projects.get(f["path"])
        owner_id = None
        owner_name = None
        if proj:
            f["project_id"] = proj["id"]
            f["project_name"] = proj["name"]
            f["specialist_name"] = proj["specialist_name"] or proj.get("user_full_name", "")
            f["original_filename"] = proj.get("original_filename", f["filename"])
            owner_id = proj.get("user_id")
            owner_name = proj.get("user_full_name") or proj.get("specialist_name")
        else:
            f["original_filename"] = f["filename"]

        # Infer owner from folder path if unset: uploads/users/<safe_name>/YYYY-MM/...
        folder = f.get("folder", "") or ""
        folder_parts = folder.split("/") if folder else []
        inferred_user_folder = None
        inferred_month = None
        if len(folder_parts) >= 2 and folder_parts[0] == "users":
            inferred_user_folder = folder_parts[1]
            if len(folder_parts) >= 3:
                inferred_month = folder_parts[2]

        if owner_id is None and inferred_user_folder:
            # Try to match inferred folder name to a user
            owner_name = owner_name or inferred_user_folder

        f["user_id"] = owner_id
        f["owner_name"] = owner_name or "Legacy"

        # Compute folder_path for admin view grouping
        if inferred_user_folder:
            safe = inferred_user_folder
            month = inferred_month or datetime.fromtimestamp(f["modified_at"]).strftime("%Y-%m")
            f["folder_path"] = f"users/{safe}/{month}"
        elif owner_name:
            safe = _safe_folder_name(owner_name, owner_id or "legacy")
            month = datetime.fromtimestamp(f["modified_at"]).strftime("%Y-%m")
            f["folder_path"] = f"users/{safe}/{month}"
        else:
            f["folder_path"] = "users/legacy"

        # Filter for non-admin: keep files only inside their own user folder OR matched by owner_id
        if not is_admin:
            in_user_folder = bool(user_safe and inferred_user_folder == user_safe)
            owns_via_db = owner_id and owner_id == current_user_id
            if not (in_user_folder or owns_via_db):
                continue
        enriched.append(f)

    enriched.sort(key=lambda x: x.get("modified_at", 0), reverse=True)
    return enriched


def get_file_owner_id(file_path: str) -> Optional[str]:
    """Return user_id owning this file (via projects.file_path match), or None."""
    conn = connect_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT user_id FROM projects WHERE file_path = ? LIMIT 1", (file_path,))
    row = cursor.fetchone()
    conn.close()
    if row:
        return dict(row).get("user_id")
    return None


def get_file_text_preview(file_path: str, max_chars: int = 2000) -> str:
    """Extract text preview from a DOCX or PDF file."""
    import os
    if not os.path.exists(file_path):
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    try:
        if ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
        elif ext == ".pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                parts = []
                for page in doc:
                    parts.append(page.get_text())
                text = "\n".join(parts)
                doc.close()
            except ImportError:
                text = "[PDF preview учун PyMuPDF талаб қилинади]"
    except Exception as e:
        text = f"[Preview хатоси: {str(e)}]"
    
    return text[:max_chars] if text else "[Бўш файл]"


# ═══════════════════════════════════════════════
# Synonyms CRUD
# ═══════════════════════════════════════════════

def save_synonym(word: str, synonym: str, lang: str, source: str = 'ai', created_by: str = ''):
    """Save a synonym to the database (upsert)."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO synonyms (word, synonym, lang, source, created_by, frequency)
        VALUES (?, ?, ?, ?, ?, 0)
        ON CONFLICT(word, synonym, lang) DO UPDATE SET
            frequency = frequency  -- keep existing frequency on conflict
    ''', (word.strip(), synonym.strip(), lang, source, created_by))
    conn.commit()
    conn.close()

def save_synonyms_batch(word: str, synonyms: list, lang: str, source: str = 'ai', created_by: str = ''):
    """Save multiple synonyms for a word at once."""
    conn = connect_db()
    cursor = conn.cursor()
    for syn in synonyms:
        if syn and syn.strip():
            cursor.execute('''
                INSERT INTO synonyms (word, synonym, lang, source, created_by, frequency)
                VALUES (?, ?, ?, ?, ?, 0)
                ON CONFLICT(word, synonym, lang) DO NOTHING
            ''', (word.strip(), syn.strip(), lang, source, created_by))
    conn.commit()
    conn.close()

def increment_synonym_frequency(word: str, synonym: str, lang: str):
    """Increment the selection frequency of a synonym."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute('''
        UPDATE synonyms SET frequency = frequency + 1
        WHERE word = ? AND synonym = ? AND lang = ?
    ''', (word.strip(), synonym.strip(), lang))
    conn.commit()
    conn.close()

def get_synonyms(word: str = None, lang: str = None, limit: int = 200):
    """Get synonyms, optionally filtered by word and lang."""
    conn = connect_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    query = "SELECT * FROM synonyms WHERE 1=1"
    params = []
    if word:
        query += " AND word LIKE ?"
        params.append(f"%{word}%")
    if lang:
        query += " AND lang = ?"
        params.append(lang)
    query += " ORDER BY frequency DESC, created_at DESC LIMIT ?"
    params.append(limit)
    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_all_synonyms(limit: int = 10000):
    """Get all synonyms for the synonyms page."""
    conn = connect_db()
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM synonyms ORDER BY frequency DESC, created_at DESC LIMIT ?", (limit,))
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]

def delete_synonym(syn_id: int):
    """Delete a synonym by ID."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM synonyms WHERE id = ?", (syn_id,))
    conn.commit()
    conn.close()

def update_synonym(syn_id: int, data: dict):
    """Update a synonym's word, synonym text, or other fields."""
    conn = connect_db()
    cursor = conn.cursor()
    fields = []
    values = []
    for k in ['word', 'synonym', 'lang', 'frequency']:
        if k in data:
            fields.append(f"{k} = ?")
            values.append(data[k])
    if fields:
        values.append(syn_id)
        cursor.execute(f"UPDATE synonyms SET {', '.join(fields)} WHERE id = ?", tuple(values))
        conn.commit()
    conn.close()

def get_synonyms_count():
    """Get total synonym count."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM synonyms")
    count = cursor.fetchone()[0]
    conn.close()
    return count


# ═══════════════════════════════════════════════
# User Profile & Password 
# ═══════════════════════════════════════════════

def update_user_profile(user_id: str, name: str = None, email: str = None):
    """Update user profile information."""
    conn = connect_db()
    cursor = conn.cursor()
    if name:
        cursor.execute("UPDATE users SET name = ? WHERE id = ?", (name, user_id))
    if email:
        cursor.execute("UPDATE users SET email = ? WHERE id = ?", (email.lower(), user_id))
    conn.commit()
    conn.close()

def set_user_password(user_id: str, hashed_password: str):
    """Set or update user password (for Google OAuth users adding password)."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("UPDATE users SET password = ? WHERE id = ?", (hashed_password, user_id))
    conn.commit()
    conn.close()

def get_user_password_hash(user_id: str):
    """Get user password hash for verification."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("SELECT password FROM users WHERE id = ?", (user_id,))
    row = cursor.fetchone()
    conn.close()
    return row[0] if row else None
def increment_synonym_frequency(word: str, synonym: str, lang: str = 'uz'):
    """Record that a specific synonym was chosen, building the Tier 1 knowledge base."""
    conn = connect_db()
    cursor = conn.cursor()
    # Update frequency
    cursor.execute('''
        UPDATE synonyms 
        SET frequency = frequency + 1 
        WHERE (word = ? AND synonym = ? AND lang = ?)
    ''', (word, synonym, lang))
    
    if cursor.rowcount == 0:
        # If not found, add it
        cursor.execute('''
            INSERT INTO synonyms (word, synonym, lang, frequency, source, author)
            VALUES (?, ?, ?, 1, 'user_choice', 'expert')
        ''', (word, synonym, lang))
        
    conn.commit()
    conn.close()

def add_synonyms_batch(word: str, synonyms: list, lang: str, author: str = 'ai'):
    """Immediate save of AI-generated synonyms to the DB."""
    conn = connect_db()
    cursor = conn.cursor()
    for syn in synonyms:
        try:
            cursor.execute('''
                INSERT OR IGNORE INTO synonyms (word, synonym, lang, source, author)
                VALUES (?, ?, ?, 'ai', ?)
            ''', (word, syn, lang, author))
        except: pass
    conn.commit()
    conn.close()

def get_synonyms_with_stats(word: str, lang: str):
    """Get synonyms for a word including frequency-based probability."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute('''
        SELECT synonym, frequency, probability_scale, part_of_speech, source, author
        FROM synonyms 
        WHERE word = ? AND lang = ?
        ORDER BY frequency DESC, probability_scale DESC
    ''', (word, lang))
    rows = cursor.fetchall()
    conn.close()
    return [dict(r) for r in rows]
    conn.close()
def update_source_lang(text_id: str, lang: str):
    """Save the selected source language for a project."""
    conn = connect_db()
    cursor = conn.cursor()
    cursor.execute("UPDATE projects SET source_lang = ? WHERE id = ?", (lang, text_id))
    conn.commit()
    conn.close()
