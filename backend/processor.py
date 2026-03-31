import docx
import re
import os
import time
from typing import List, Dict, Any
from anthropic import Anthropic
from dotenv import load_dotenv
import json
from langdetect import detect, DetectorFactory
DetectorFactory.seed = 0 # Consistent detection

load_dotenv()

IMGS_DIR = os.path.join("temp_files", "imgs")
os.makedirs(IMGS_DIR, exist_ok=True)


class ParagraphAligner:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = docx.Document(file_path)
        self.table = self.doc.tables[0] if self.doc.tables else None
        api_key = os.getenv("ANTHROPIC_API_KEY")
        self.client = Anthropic(api_key=api_key) if api_key else None

    # ──────────────────────────────────────────────
    # Image extraction
    # ──────────────────────────────────────────────
    def extract_image_url(self, paragraph) -> str:
        """Extract image from paragraph, save to disk, return §IMG:url§ token."""
        try:
            from docx.oxml.ns import qn
            blips = paragraph._p.findall('.//' + qn('a:blip'))
            for blip in blips:
                rId = blip.get(qn('r:embed'))
                if rId:
                    try:
                        img_part = paragraph.part.related_parts[rId]
                        raw_ext = img_part.content_type.split('/')[-1]
                        ext = 'jpg' if 'jpeg' in raw_ext else raw_ext
                        fname = f"img_{abs(hash(rId + self.file_path))}.{ext}"
                        fpath = os.path.join(IMGS_DIR, fname)
                        if not os.path.exists(fpath):
                            with open(fpath, 'wb') as f:
                                f.write(img_part.blob)
                        return f'§IMG:http://localhost:8000/static/imgs/{fname}§'
                    except Exception as e:
                        print(f"Image extract error: {e}")
        except Exception as e:
            print(f"Image scan error: {e}")
        return '§IMG§'

    # ──────────────────────────────────────────────
    # Nested table → HTML
    # ──────────────────────────────────────────────
    def table_to_html(self, tbl) -> str:
        html = '<table style="border-collapse:collapse;width:100%;font-size:0.78rem;margin:4px 0">'
        for row in tbl.rows:
            html += '<tr>'
            for cell in row.cells:
                txt = cell.text.replace('\n', '<br/>')
                html += f'<td style="border:1px solid #ccc;padding:3px 6px">{txt}</td>'
            html += '</tr>'
        html += '</table>'
        return f'§TBL:{html}§'

    # ──────────────────────────────────────────────
    # Get all items from a cell (paragraphs + tables in order)
    # ──────────────────────────────────────────────
    def get_cell_items(self, cell) -> List[str]:
        """
        Iterates through cell XML children preserving order.
        Returns list of strings:
          - plain text for text paragraphs
          - §IMG:url§ for images
          - §TBL:<html>§ for nested tables
          - §MAT:xml§ for mathematical formulas (OMath)
        """
        items = []
        from docx.oxml.ns import qn
        OMATH = '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath'
        OMATH_PARA = '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMathPara'

        for child in cell._tc:
            local = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if local == 'p':
                from docx.text.paragraph import Paragraph
                p = Paragraph(child, cell)
                text = p.text.strip()
                
                # Check for OMath within paragraph
                maths = child.findall('.//' + OMATH)
                math_paras = child.findall('.//' + OMATH_PARA)
                
                has_drawing = 'w:drawing' in child.xml or 'graphicData' in child.xml

                if math_paras or maths:
                    # Formula placeholder (keeping it simple for tokenization)
                    items.append(f'§MAT:Formula {len(items)+1}§')
                elif has_drawing:
                    items.append(self.extract_image_url(p))
                elif text:
                    items.append(text)

            elif local == 'tbl':
                from docx.table import Table
                try:
                    tbl = Table(child, cell)
                    items.append(self.table_to_html(tbl))
                except Exception:
                    items.append('§TBL:<em>[Жадвал]</em>§')
            
            elif local in ['oMath', 'oMathPara']:
                items.append(f'§MAT:Formula {len(items)+1}§')

        return items

    # ──────────────────────────────────────────────
    # Split cell items into blocks (by section markers)
    # ──────────────────────────────────────────────
    def get_blocks(self, cell) -> List[Dict[str, Any]]:
        items = self.get_cell_items(cell)
        blocks = []
        current_block = {"marker": None, "paragraphs": []}

        for text in items:
            if not text:
                continue

            # Special §...§ items: Tables and Formulas are anchors, images are not.
            is_anchor_token = text.startswith('§TBL:') or text.startswith('§MAT:')
            is_anchor = False
            
            if is_anchor_token:
                is_anchor = True
            elif not text.startswith('§'):
                # Regular text anchors like {1}, 〈1〉, 1.1 etc.
                if (text.startswith('{') and text.endswith('}')) or \
                   re.match(r'^〈\s*\d+\s*〉$', text) or \
                   re.match(r'^(\d+\.)+\d*$', text) or \
                   re.match(r'^\d+$', text):
                    is_anchor = True

            if is_anchor:
                # Save previous block if it has anything
                if current_block["marker"] or current_block["paragraphs"]:
                    blocks.append(current_block)
                
                # Start new block. If it's a table/math, §TBL: or §MAT: itself is the marker.
                marker_text = text
                if is_anchor_token:
                    # Keep raw token as marker
                    marker_text = text 
                
                current_block = {"marker": marker_text, "paragraphs": []}
            else:
                current_block["paragraphs"].append(text)

        if current_block["marker"] or current_block["paragraphs"]:
            blocks.append(current_block)
        return blocks

    # ──────────────────────────────────────────────
    # Jaccard similarity (word-level)
    # ──────────────────────────────────────────────
    def get_similarity(self, s1: str, s2: str) -> float:
        if not s1 or not s2: return 0.0
        if s1.startswith('§') or s2.startswith('§'): return 0.0
        w1 = set(re.sub(r'[^\w\s]', '', s1.lower()).split())
        w2 = set(re.sub(r'[^\w\s]', '', s2.lower()).split())
        if not w1 or not w2: return 0.0
        return len(w1 & w2) / max(len(w1), len(w2))

    # ──────────────────────────────────────────────
    # Paragraph-level alignment (NOT sentence-level)
    # This prevents the "wrong row" bug
    # ──────────────────────────────────────────────
    def align_paragraphs(self, en_items: List[str], ru_items: List[str], uz_items: List[str]) -> List[Dict[str, Any]]:
        """
        Aggressive 3-way alignment. 
        Each EN paragraph is the master. We find best matching RU and UZ segments.
        If RU/UZ has significantly more paragraphs, we append them to the nearest matching block
        to maintain overall structure.
        """
        def find_best(en: str, candidates: List[str], used: set, center: int) -> int:
            if not candidates: return -1
            # Special items (images/tables/formulas) match strictly by position
            if en.startswith('§'):
                window = 3
                for j in range(max(0, center - window), min(len(candidates), center + window + 1)):
                    if j not in used and candidates[j].startswith('§'):
                        return j
                # Fallback to nearest unused special or just nearest unused
                for j in range(len(candidates)):
                    if j not in used: return j
                return -1

            en_words = len(en.split())
            
            # SHORT HEADINGS (1-4 words like "INTRODUCTION", "1.1 SCOPE"):
            # Use strict positional matching — don't try Jaccard on these
            if en_words <= 4:
                # First, try to find a short candidate near the expected position  
                for j in range(max(0, center - 2), min(len(candidates), center + 3)):
                    if j not in used:
                        cand_words = len(candidates[j].split())
                        # Match short with short (heading with heading)
                        if cand_words <= 6:
                            return j
                # If no short candidate nearby, just take the positional match
                if center < len(candidates) and center not in used:
                    return center
                for j in range(len(candidates)):
                    if j not in used: return j
                return -1

            # Normal text: use Jaccard + distance penalty
            best_idx, max_score = -1, -1.0
            search_range = 5 # Broad search
            for j in range(max(0, center - search_range), min(len(candidates), center + search_range + 1)):
                if j in used: continue
                similarity = self.get_similarity(en, candidates[j])
                dist_penalty = 1.0 - (abs(j - center) / len(candidates)) if len(candidates) > 0 else 1.0
                score = similarity * 0.7 + dist_penalty * 0.3
                if score > max_score:
                    max_score, best_idx = score, j
            
            if best_idx == -1:
                # Take next available
                for j in range(len(candidates)):
                    if j not in used: return j
            return best_idx

        n_en = len(en_items)
        if n_en == 0:
            # No English, but maybe translations? (Unlikely in this doc structure)
            # Just group RU/UZ into one block if they exist
            res = []
            if ru_items or uz_items:
                res.append({
                    "en": "",
                    "ru_v1": " ".join(ru_items),
                    "uz_v1": " ".join(uz_items)
                })
            return res

        res = []
        used_ru, used_uz = set(), set()
        
        for i, en in enumerate(en_items):
            ru_center = round(i * len(ru_items) / n_en) if len(ru_items) > 0 else 0
            uz_center = round(i * len(uz_items) / n_en) if len(uz_items) > 0 else 0
            
            ri = find_best(en, ru_items, used_ru, ru_center)
            ui = find_best(en, uz_items, used_uz, uz_center)
            
            if ri != -1: used_ru.add(ri)
            if ui != -1: used_uz.add(ui)
            
            res.append({
                "en": en,
                "ru_v1": ru_items[ri] if ri != -1 else "",
                "uz_v1": uz_items[ui] if ui != -1 else ""
            })

        # Final cleanup: Append all unused RU/UZ items to the LAST row of the block
        # to ensure no data is lost.
        if res:
            leftover_ru = [ru_items[j] for j in range(len(ru_items)) if j not in used_ru]
            leftover_uz = [uz_items[j] for j in range(len(uz_items)) if j not in used_uz]
            if leftover_ru:
                res[-1]["ru_v1"] = (res[-1]["ru_v1"] + "\n\n" + "\n\n".join(leftover_ru)).strip()
            if leftover_uz:
                res[-1]["uz_v1"] = (res[-1]["uz_v1"] + "\n\n" + "\n\n".join(leftover_uz)).strip()
        
        # ──────────────────────────────────────────────
        # NEW: Universal Sync for technical content (IMG/TBL/MAT)
        # ──────────────────────────────────────────────
        for row in res:
            # Collect all special tokens found in this row across all cells
            all_txt = (row['en'] + " " + row['ru_v1'] + " " + row['uz_v1'])
            specials = re.findall(r'§(?:IMG|TBL|MAT):[^§]*§|§IMG§', all_txt)
            for s in specials:
                # If row has a special but it is missing in any of the columns, 
                # we broadcast it (assuming it is shared technical context)
                for col in ['en', 'ru_v1', 'uz_v1']:
                    if s not in row[col]:
                        # If cell is EMPTY, just put it. If not, append it?
                        # In pharma docs, symbols are usually standalone or attached.
                        # We append it with a newline if the cell has text.
                        if not row[col].strip():
                            row[col] = s
                        else:
                            row[col] = row[col] + "\n" + s
        
        return res

    # ──────────────────────────────────────────────
    # Detect language of a text block
    # ──────────────────────────────────────────────
    def detect_language(self, text: str) -> str:
        """Detect if text is English, Russian, or Uzbek. Returns 'en', 'ru', or 'uz'."""
        if not text or len(text.strip()) < 3:
            return 'en'
        try:
            lang = detect(text)
            # langdetect returns 'en', 'ru', 'uk', 'tr', etc.
            if lang == 'ru' or lang == 'uk' or lang == 'bg':
                return 'ru'
            elif lang == 'tr' or lang == 'uz' or lang == 'az':
                return 'uz'
            else:
                return 'en'
        except Exception:
            # Fallback: check for Cyrillic characters
            cyrillic_count = sum(1 for c in text if '\u0400' <= c <= '\u04FF')
            latin_count = sum(1 for c in text if 'a' <= c.lower() <= 'z')
            if cyrillic_count > latin_count:
                # Distinguish RU from UZ-Cyrillic by checking for Uzbek-specific letters
                uz_chars = sum(1 for c in text if c in 'ҳҚқҒғЎўҲ')
                if uz_chars > 0:
                    return 'uz'
                return 'ru'
            return 'en'

    # ──────────────────────────────────────────────
    # Extract all text from document (paragraphs + tables)
    # ──────────────────────────────────────────────
    def extract_all_text(self) -> List[str]:
        """Extract all text paragraphs from the document body."""
        paragraphs = []
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract from any tables (single-column or otherwise)
        if self.table:
            for row in self.table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in paragraphs:
                        paragraphs.append(text)
        return paragraphs

    # ──────────────────────────────────────────────
    # Process single-language document (no 3-column table)
    # ──────────────────────────────────────────────
    def process_single_language(self) -> List[Dict[str, Any]]:
        """Process a document that is in a single language without column separation."""
        text_id = os.path.basename(self.file_path).replace(".docx", "")
        paragraphs = self.extract_all_text()
        
        if not paragraphs:
            raise ValueError("Ҳужжатда матн топилмади / No text found in document")
        
        # Detect language from combined text
        combined = " ".join(paragraphs[:10])  # Use first 10 paragraphs for detection
        lang = self.detect_language(combined)
        
        aligned_data = []
        for idx, para_text in enumerate(paragraphs):
            row = {
                "type": "content",
                "en": "",
                "ru_v1": "",
                "ru_proposed": "",
                "uz_v1": "",
                "uz_proposed": "",
                "status": "review",
                "sentence_no": idx + 1,
                "display_no": str(idx + 1),
                "text_id": text_id,
                "notes": ""
            }
            # Map text to the detected language column
            if lang == 'en':
                row["en"] = para_text
            elif lang == 'ru':
                row["ru_v1"] = para_text
                row["ru_proposed"] = para_text
            else:  # uz
                row["uz_v1"] = para_text
                row["uz_proposed"] = para_text
            
            aligned_data.append(row)
        
        return aligned_data

    # ──────────────────────────────────────────────
    # Check if document has a valid 3-column table
    # ──────────────────────────────────────────────
    def has_three_column_table(self) -> bool:
        """Check if the document contains a table with at least 3 columns."""
        if not self.table:
            return False
        for row in self.table.rows:
            if len(row.cells) >= 3:
                return True
        return False

    # ──────────────────────────────────────────────
    # Process "Ready Form" (Skip alignment, 1 row = 1 row)
    # ──────────────────────────────────────────────
    def process_ready_form(self) -> List[Dict[str, Any]]:
        # If no 3-column table, fall back to single-language processing
        if not self.has_three_column_table():
            return self.process_single_language()
        
        aligned_data = []
        text_id = os.path.basename(self.file_path).replace(".docx", "")
        
        # Start from row 0 if no clear header, but usually row 1 is data
        start_idx = 0
        if len(self.table.rows) > 1:
            # Check if first row is just headers
            cell0 = self.table.rows[0].cells[0].text.lower()
            if 'english' in cell0 or 'en' == cell0.strip():
                start_idx = 1

        for idx, row in enumerate(self.table.rows[start_idx:]):
            if len(row.cells) < 3: continue
            
            en_items = self.get_cell_items(row.cells[0])
            ru_items = self.get_cell_items(row.cells[1])
            uz_items = self.get_cell_items(row.cells[2])
            
            # Combine multiple items in one cell with newlines for editing
            en_txt = "\n\n".join(en_items)
            ru_txt = "\n\n".join(ru_items)
            uz_txt = "\n\n".join(uz_items)
            
            if not en_txt and not ru_txt and not uz_txt:
                continue

            aligned_data.append({
                "type": "content",
                "en": en_txt,
                "ru_v1": ru_txt,
                "ru_proposed": ru_txt, # Direct copy for ready form
                "uz_v1": uz_txt,
                "uz_proposed": uz_txt, # Direct copy for ready form
                "status": "aligned",
                "sentence_no": idx + 1,
                "display_no": str(idx + 1),
                "text_id": text_id,
                "notes": ""
            })
            
        return aligned_data

    # ──────────────────────────────────────────────
    # Main process: build aligned data (AUTO alignment)
    # ──────────────────────────────────────────────
    def process(self) -> List[Dict[str, Any]]:
        # If no 3-column table, fall back to single-language processing
        if not self.has_three_column_table():
            return self.process_single_language()

        # Count actual content rows (3+ cells)
        content_rows = [r for r in self.table.rows if len(r.cells) >= 3]
        
        # If multiple rows with 3+ cells, this is a multi-row table
        # => each row is a separate content unit => use ready_form logic
        if len(content_rows) > 1:
            return self.process_ready_form()

        # Single content row: all text for each language is in one big cell
        content_row = content_rows[0] if content_rows else None
        if not content_row:
            return self.process_single_language()

        en_blocks = self.get_blocks(content_row.cells[0])
        ru_blocks = self.get_blocks(content_row.cells[1])
        uz_blocks = self.get_blocks(content_row.cells[2])

        aligned_data = []
        sentence_counter = 1
        text_id = os.path.basename(self.file_path).replace(".docx", "")
        max_blocks = max(len(en_blocks), len(ru_blocks), len(uz_blocks))

        for i in range(max_blocks):
            en_b = en_blocks[i] if i < len(en_blocks) else {"marker": None, "paragraphs": []}
            ru_b = ru_blocks[i] if i < len(ru_blocks) else {"marker": None, "paragraphs": []}
            uz_b = uz_blocks[i] if i < len(uz_blocks) else {"marker": None, "paragraphs": []}

            # Marker row
            aligned_data.append({
                "type": "marker",
                "en": en_b["marker"] or "",
                "ru_v1": ru_b["marker"] or "",
                "uz_v1": uz_b["marker"] or "",
                "ru_proposed": ru_b["marker"] or "",
                "uz_proposed": uz_b["marker"] or "",
                "status": "aligned",
                "sentence_no": 0,
                "display_no": "",
                "text_id": text_id,
                "notes": ""
            })

            # Content rows — aligned at PARAGRAPH level
            block_rows = self.align_paragraphs(
                en_b["paragraphs"], ru_b["paragraphs"], uz_b["paragraphs"]
            )

            for r in block_rows:
                aligned_data.append({
                    "type": "content",
                    "en":          r.get("en", ""),
                    "ru_v1":       r.get("ru_v1", ""),
                    "uz_v1":       r.get("uz_v1", ""),
                    "ru_proposed": r.get("ru_v1", ""),
                    "uz_proposed": r.get("uz_v1", ""),
                    "status": "review",
                    "sentence_no": sentence_counter,
                    "display_no":  str(sentence_counter),
                    "text_id":     text_id,
                    "notes":       ""
                })
                sentence_counter += 1

        return aligned_data


# ──────────────────────────────────────────────────
# DOCX Export (Logica Expert Style)
# ──────────────────────────────────────────────────
from docx.shared import Pt, RGBColor, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """Set background color of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def parse_html_table(html: str) -> List[List[str]]:
    """Simple parser for the §TBL:<html>§ format."""
    rows = []
    # Match <tr> items
    tr_matches = re.findall(r'<tr>(.*?)</tr>', html, re.DOTALL)
    for tr in tr_matches:
        # Match <td> items
        cells = re.findall(r'<td[^>]*>(.*?)</td>', tr, re.DOTALL)
        # Clean <br/> and other tags
        clean_cells = [re.sub(r'<[^>]*>', ' ', c).strip() for c in cells]
        rows.append(clean_cells)
    return rows

def export_to_docx(data: List[Dict[str, Any]], output_path: str):
    doc = docx.Document()
    
    # 1. Set to A4 Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # 2. Add Title
    title_p = doc.add_paragraph()
    title_p.alignment = 1 # Center
    title_run = title_p.add_run("PHARMA TRANSLATION REPORT")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Arial'
    
    file_p = doc.add_paragraph()
    file_p.alignment = 1 # Center
    file_run = file_p.add_run(f"Source: {os.path.basename(output_path).replace('confirmed_', '')}")
    file_run.font.size = Pt(10)
    file_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_paragraph() # Spacer

    # 3. Main Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set total width approx 10 inches (A4 landscape width is ~11.7)
    # 3 columns = ~3.33 inches each
    col_w = Inches(3.4)
    
    # 4. Header Row (Blue #1E293B, White Bold Text)
    hdr_cells = table.rows[0].cells
    headers = ['ENGLISH (Original)', 'RUSSIAN (Confirmed)', 'UZBEK (Confirmed)']
    for i, h_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = col_w
        set_cell_shading(cell, '1E293B')
        para = cell.paragraphs[0]
        para.alignment = 1 # Center
        run = para.add_run(h_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Arial'

    # 5. Content Rows
    for idx, item in enumerate(data):
        row = table.add_row()
        row_cells = row.cells
        
        is_marker = item.get("type") == "marker"
        
        # Color logic: Markers are light gray, rows are white/off-white
        bg_color = 'E2E8F0' if is_marker else ('FFFFFF' if idx % 2 == 0 else 'F8F9FA')
        
        # Mapping for languages (Logic: use Proposed if it exists and differs from V1, otherwise V1)
        # However, for simplicity and ensuring "Latest changes", we use Proposed if not empty.
        lang_map = {
            0: item.get('en', ''),
            1: item.get('ru_proposed') or item.get('ru_v1', ''),
            2: item.get('uz_proposed') or item.get('uz_v1', '')
        }
        
        for li in range(3):
            cell = row_cells[li]
            cell.width = col_w
            set_cell_shading(cell, bg_color)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            raw_text = lang_map[li]
            # Clear existing paragraph if any
            while len(cell.paragraphs) > 0:
                p = cell.paragraphs[0]
                p._element.getparent().remove(p._element)
            
            # Split by tokens §...§
            parts = re.split(r'(§(?:IMG|TBL):[^§]*§|§IMG§)', str(raw_text))
            for part in parts:
                if not part: continue
                
                if part.startswith('§TBL:'):
                    html = part[5:-1]
                    try:
                        inner_data = parse_html_table(html)
                        if inner_data:
                            inner_table = cell.add_table(rows=len(inner_data), cols=len(inner_data[0]))
                            inner_table.style = 'Table Grid'
                            for ri, r_data in enumerate(inner_data):
                                for ci, val in enumerate(r_data):
                                    if ci < len(inner_table.rows[ri].cells):
                                        inner_table.rows[ri].cells[ci].text = val
                                        for p in inner_table.rows[ri].cells[ci].paragraphs:
                                            for run in p.runs: run.font.size = Pt(8)
                    except Exception as e:
                        cell.add_paragraph(f"[Table Error: {e}]")
                
                elif part.startswith('§MAT:'):
                    p = cell.add_paragraph("[ФОРМУЛА]") # Localized placeholder
                    p.alignment = 1 # Center
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                
                elif part.startswith('§IMG:'):
                    match = re.search(r'/imgs/([^§/]+)$', part[:-1])
                    if match:
                        fname = match.group(1)
                        fpath = os.path.join(IMGS_DIR, fname)
                        if os.path.exists(fpath):
                            try:
                                p = cell.add_paragraph()
                                p.alignment = 1
                                r = p.add_run()
                                r.add_picture(fpath, width=Inches(2.8))
                            except Exception as e:
                                cell.add_paragraph(f"[Image Error: {e}]")
                        else:
                            cell.add_paragraph(f"[Image Missing: {fname}]")
                    else:
                        cell.add_paragraph("[РАСМ]")
                elif part == '§IMG§':
                    cell.add_paragraph("[РАСМ]")
                else:
                    # Regular text
                    p = cell.add_paragraph(part)
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'
                        if is_marker:
                            run.bold = True
                        if li == 0: # English italicize slightly? Or keep clean.
                            # run.italic = True
                            pass

    doc.save(output_path)
