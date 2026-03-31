import docx
import os

def analyze_docx(file_path):
    doc = docx.Document(file_path)
    print(f"Analyzing: {file_path}")
    print(f"Number of tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}:")
        print(f"Number of rows: {len(table.rows)}")
        print(f"Number of columns: {len(table.columns)}")
        
        # Look at the first few rows to understand structure
        for row_idx, row in enumerate(table.rows[:5]):
            print(f"Row {row_idx}:")
            for col_idx, cell in enumerate(row.cells):
                # Count paragraphs in each cell
                p_count = len(cell.paragraphs)
                text_preview = cell.text[:50].replace('\n', ' ')
                print(f"  Col {col_idx}: {p_count} paragraphs. Preview: {text_preview}...")

if __name__ == "__main__":
    file_path = r"c:\Users\Администратор\Desktop\2\1023. рус авто v2.docx"
    if os.path.exists(file_path):
        analyze_docx(file_path)
    else:
        print(f"File not found: {file_path}")
