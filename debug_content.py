import docx

def debug_alignment(file_path):
    doc = docx.Document(file_path)
    table = doc.tables[0]
    row = table.rows[1]
    
    en = [p.text.strip() for p in row.cells[0].paragraphs if p.text.strip()]
    ru = [p.text.strip() for p in row.cells[1].paragraphs if p.text.strip()]
    uz = [p.text.strip() for p in row.cells[2].paragraphs if p.text.strip()]
    
    print("--- ENGLISH (First 10) ---")
    for i, t in enumerate(en[:10]): print(f"{i}: {t[:100]}")
    
    print("\n--- RUSSIAN (First 10) ---")
    for i, t in enumerate(ru[:10]): print(f"{i}: {t[:100]}")
    
    print("\n--- UZBEK (First 10) ---")
    for i, t in enumerate(uz[:10]): print(f"{i}: {t[:100]}")

if __name__ == "__main__":
    debug_alignment(r"c:\Users\Администратор\Desktop\2\1023. рус авто v2.docx")
